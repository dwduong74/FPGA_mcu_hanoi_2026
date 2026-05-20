"""
gen_code_explain.py
Tạo: Tài_liệu_giải_thích_code.docx
Giải thích chi tiết từng dòng code dự án Digital Clock SN32F407
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()
for sec in doc.sections:
    sec.top_margin = Cm(2); sec.bottom_margin = Cm(2)
    sec.left_margin = Cm(2.5); sec.right_margin = Cm(2)

def h1(t): doc.add_heading(t, 1)
def h2(t): doc.add_heading(t, 2)
def h3(t): doc.add_heading(t, 3)
def h4(t): doc.add_heading(t, 4)
def p(t):  doc.add_paragraph(t)
def li(t): doc.add_paragraph(t, style='List Bullet')
def li2(t):doc.add_paragraph(t, style='List Bullet 2')
def blank():doc.add_paragraph()

def code(t):
    para = doc.add_paragraph()
    run = para.add_run(t)
    run.font.name = 'Courier New'
    run.font.size = Pt(8.5)
    para.paragraph_format.left_indent = Cm(0.8)
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after  = Pt(2)
    return para

def note(t):
    para = doc.add_paragraph('► ' + t)
    for run in para.runs:
        run.italic = True
        run.font.color.rgb = RGBColor(0xCC,0x44,0x00)

def tbl(headers, rows):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'
    for i,h in enumerate(t.rows[0].cells):
        h.text = headers[i]
        for r in h.paragraphs[0].runs: r.bold = True
    for ri,row in enumerate(rows):
        for ci,val in enumerate(row):
            t.rows[ri+1].cells[ci].text = str(val)
    doc.add_paragraph()

# ─── TIÊU ĐỀ ───────────────────────────────────────────────────────────────
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run('TÀI LIỆU GIẢI THÍCH CODE CHI TIẾT\nDigital Clock — SN32F407_EVK')
r.bold = True; r.font.size = Pt(17)
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.add_run('Đọc hiểu từng dòng code • Giải thích thanh ghi • Lý do thiết kế').italic = True
blank()

p('Tài liệu này đi qua từng file source theo thứ tự thực thi, '
  'giải thích ý nghĩa của mỗi hằng số, biến, hàm, và từng dòng code quan trọng. '
  'Mục tiêu: sau khi đọc xong có thể trả lời bất kỳ câu hỏi nào về code này.')
blank()

# ═══════════════════════════════════════════════════════════════════════════
h1('PHẦN 1 — CẤU TRÚC PROJECT VÀ THỨ TỰ BIÊN DỊCH')
# ═══════════════════════════════════════════════════════════════════════════
tbl(['File', 'Vai trò', 'Được gọi từ'],
[
 ['startup_SN32F400.s','Assembly: khởi tạo stack, copy .data, zero .bss, gọi main()','Reset Vector (tự động)'],
 ['system_SN32F400.c', 'SystemInit(): cấu hình HCLK=12MHz IHRC',                   'startup, trước main()'],
 ['main.c',            'Logic chính: state machine, điều khiển đồng hồ',            'startup → main()'],
 ['CT16B1.c',          'Timer 1ms: CT16B1_Init(), CT16B1_IRQHandler()',              'main.c'],
 ['CT16B0.c',          'Buzzer PWM: CT16B0_Init(), ISR clear cờ',                   'main.c'],
 ['GPIO.c',            'GPIO_Init(): cấu hình tất cả chân',                         'main.c'],
 ['PFPA.c',            'PFPA_Init(): ánh xạ chân ngoại vi',                         'main.c'],
 ['Segment.c',         'SEGMENT_TABLE[], Digital_Scan(), segment_buff[]',            'main.c'],
 ['KeyScan.c',         'KeyScan(): quét ma trận phím, debounce 50ms',               'main.c'],
 ['EEPROM.c',          'eeprom_read(), eeprom_write() qua I2C',                     'main.c'],
 ['I2C0.c',            'I2C0_Init(), I2C_write_byte(), I2C_read_byte()',             'EEPROM.c'],
 ['WDT.c',             'WDT_Init()',                                                 'main.c'],
])

# ═══════════════════════════════════════════════════════════════════════════
h1('PHẦN 2 — main.c: HẰNG SỐ VÀ BIẾN TOÀN CỤC')
# ═══════════════════════════════════════════════════════════════════════════

h2('2.1 Hằng số trạng thái (State Constants)')
code(
'#define STATE_NORMAL          0   // Đồng hồ chạy bình thường\n'
'#define STATE_SET_HOUR        1   // Đang chỉnh giờ hiện tại\n'
'#define STATE_SET_MIN         2   // Đang chỉnh phút hiện tại\n'
'#define STATE_SET_ALARM_HOUR  3   // Đang chỉnh giờ báo thức\n'
'#define STATE_SET_ALARM_MIN   4   // Đang chỉnh phút báo thức'
)
p('Dùng số nguyên (0–4) thay vì enum để dễ quan sát trong Watch Window Keil. '
  'g_state = 0 → NORMAL; g_state = 1 → đang nháy giờ; v.v. '
  'Biến g_state điều khiển toàn bộ hành vi của chương trình — '
  'cùng một phím có tác dụng khác nhau tùy g_state.')

h2('2.2 Hằng số timing')
code(
'#define BLINK_HALF_MS   500U   // Nửa chu kỳ nháy = 500ms → ON 500ms + OFF 500ms = 1s\n'
'#define TIMEOUT_MS    30000U   // Timeout không nhấn phím = 30s = 30000ms\n'
'#define PIP_MS          300U   // Pip khi nhấn phím = 0.3s\n'
'#define ALARM_HALF_MS   500U   // Nửa chu kỳ pip-pip = 0.5s\n'
'#define ALARM_TOTAL_MS 5000U   // Tổng thời gian báo thức = 5s\n'
'#define BUZZER_PERIOD  20000   // HCLK(12MHz) / 20000 = 600Hz\n'
'#define HCLK_HZ    12000000UL  // Tần số MCU = 12 MHz'
)
note('Tất cả đơn vị là MILI-GIÂY vì bộ đếm chính (g_ms_cnt) tăng mỗi 1ms. '
     'BUZZER_PERIOD tính tần số: 12,000,000 / 20,000 = 600 Hz.')

h2('2.3 Hằng số EEPROM')
code(
'#define EEPROM_WRITE_ADDR  0xA0   // Địa chỉ I2C write: 1010 000 0\n'
'#define EEPROM_READ_ADDR   0xA1   // Địa chỉ I2C read:  1010 000 1\n'
'#define EEPROM_AL_HOUR     0x00   // Ô nhớ EEPROM lưu alarm_hour\n'
'#define EEPROM_AL_MIN      0x01   // Ô nhớ EEPROM lưu alarm_min'
)
p('AT24C02 có địa chỉ I2C 7-bit = 0x50 (mặc định, chân A2A1A0 = GND). '
  'Bit 0 của byte địa chỉ trên bus I2C: 0 = write, 1 = read. '
  'Nên byte gửi trên bus = 0xA0 (write) hoặc 0xA1 (read).')

h2('2.4 Macro LED')
code(
'#define LED_ALARM_ON()   (SN_GPIO3->BCLR = (1u << 8))  // Kéo P3.8 = 0 → LED sáng\n'
'#define LED_ALARM_OFF()  (SN_GPIO3->BSET = (1u << 8))  // Kéo P3.8 = 1 → LED tắt'
)
p('BCLR (Bit Clear Register): ghi 1 vào bit nào → chân đó xuống 0. '
  'BSET (Bit Set Register): ghi 1 vào bit nào → chân đó lên 1. '
  'Dùng BCLR/BSET thay vì ghi DATA vì BCLR/BSET là thao tác atomic — '
  'không bị race condition với ISR.')

h2('2.5 Biến toàn cục và ý nghĩa')
tbl(['Biến','Kiểu','Phạm vi','Ý nghĩa'],
[
 ['g_state',         'uint8_t', 'static', 'Trạng thái hiện tại (0–4)'],
 ['g_hour',          'uint8_t', 'static', 'Giờ hiện tại (0–23)'],
 ['g_min',           'uint8_t', 'static', 'Phút hiện tại (0–59)'],
 ['g_sec',           'uint8_t', 'static', 'Giây hiện tại (0–59)'],
 ['g_alarm_hour',    'uint8_t', 'static', 'Giờ báo thức (0–23), lưu trong EEPROM'],
 ['g_alarm_min',     'uint8_t', 'static', 'Phút báo thức (0–59), lưu trong EEPROM'],
 ['g_ms_cnt',        'uint16_t','static', 'Đếm ms trong giây hiện tại (0–999)'],
 ['g_blink_cnt',     'uint16_t','static', 'Đếm ms cho chu kỳ nháy (0–499)'],
 ['g_timeout_cnt',   'uint16_t','static', 'Đếm ms không nhấn phím (0–29999)'],
 ['g_blink_on',      'uint8_t', 'static', '1=chữ số sáng, 0=chữ số tắt (nháy)'],
 ['g_pip_ms',        'uint16_t','static', 'ms còn lại của pip ngắn (0=không pip)'],
 ['g_alarm_active',  'uint8_t', 'static', '1=đang báo thức, 0=bình thường'],
 ['g_alarm_remain_ms','uint16_t','static','ms còn lại của chuỗi báo thức (0–5000)'],
 ['g_alarm_phase_ms','uint16_t','static', 'ms còn lại trong pha ON/OFF hiện tại'],
 ['g_alarm_phase_on','uint8_t', 'static', '1=pha buzzer ON, 0=pha buzzer OFF'],
 ['timer_1ms_flag',  'uint8_t', 'extern volatile','Cờ set bởi ISR mỗi 1ms'],
 ['segment_buff[4]', 'uint8_t[]','extern','Mã LED 7 đoạn cho 4 chữ số HH.MM'],
])
note('Tất cả biến là "static" → chỉ visible trong main.c, không xung đột tên '
     'với các file khác. timer_1ms_flag là volatile vì được ghi từ ISR.')

# ═══════════════════════════════════════════════════════════════════════════
h1('PHẦN 3 — main.c: HÀM main() — TRÌNH TỰ KHỞI ĐỘNG')
# ═══════════════════════════════════════════════════════════════════════════

h2('3.1 Chuỗi init và lý do thứ tự')
code(
'int main(void)\n'
'{\n'
'    SystemInit();              // 1. Cấu hình HCLK = 12MHz IHRC\n'
'    SystemCoreClockUpdate();   // 2. Cập nhật biến SystemCoreClock\n'
'    PFPA_Init();               // 3. Ánh xạ chân: I2C0→P0.10/P0.11, CT16B0→P3.0\n'
'    NotPinOut_GPIO_init();     // 4. Pull-up chân không dùng (chống nhiễu)\n'
'    SN_SYS0->EXRSTCTRL_b.RESETDIS = 0;  // 5. Cho phép chân RESET ngoài\n'
'    GPIO_Init();               // 6. Cấu hình GPIO: segment, COM, phím, LED\n'
'    WDT_Init();                // 7. Bật Watchdog (PHẢI trước khi bắt đầu vòng lặp)\n'
'    CT16B0_Init();             // 8. Buzzer PWM (timer chạy ngay)\n'
'    SN_PFPA->CT16B0_b.PWM0=1; // 9. Nối PWM0 ra chân P3.0 (sau init mới được)\n'
'    buzzer_off();              // 10. Tắt buzzer ngay (MR0=0)\n'
'    CT16B1_Init();             // 11. Timer 1ms + bật NVIC IRQ\n'
'    I2C0_Init();               // 12. I2C để đọc EEPROM\n'
'    load_alarm_from_eeprom();  // 13. Đọc giờ báo thức từ EEPROM\n'
'    g_blink_on = 1;\n'
'    update_display();          // 14. Hiển thị 00:00 ngay lập tức\n'
'    while (1) { ... }\n'
'}'
)
tbl(['Bước','Lý do thứ tự quan trọng'],
[
 ['1 trước tất cả','SystemInit() cần chạy trước khi dùng bất kỳ ngoại vi nào (clock chưa đúng = timing sai)'],
 ['3 trước 8,12','PFPA phải gán chân trước khi init CT16B0 và I2C0, không thì chân không được kết nối'],
 ['7 sớm','WDT_Init() phải gọi sớm để bắt đầu đếm; sau đó feed trong while(1)'],
 ['11 cuối cùng','CT16B1 (ngắt 1ms) bật sau cùng để không có ISR fire vào khi các biến chưa init'],
 ['13 trước 14','Phải đọc EEPROM trước khi update_display() hiển thị giờ báo thức'],
])

# ═══════════════════════════════════════════════════════════════════════════
h1('PHẦN 4 — main.c: VÒNG LẶP CHÍNH while(1)')
# ═══════════════════════════════════════════════════════════════════════════

h2('4.1 Cấu trúc tổng thể của vòng lặp')
code(
'while (1)\n'
'{\n'
'    __WDT_FEED_VALUE;              // A. Feed watchdog\n'
'\n'
'    if (timer_1ms_flag)            // B. Khối 1ms (do ISR set cờ)\n'
'    {\n'
'        timer_1ms_flag = 0;\n'
'        Digital_Scan();            //  B1. Quét LED\n'
'        g_ms_cnt++;                //  B2. Đếm ms\n'
'        if (g_ms_cnt >= 1000) { tick_second(); } //  B3. 1 giây\n'
'        if (g_state != STATE_NORMAL) {            //  B4. Blink + timeout\n'
'            g_blink_cnt++; ... g_timeout_cnt++; ...\n'
'        }\n'
'        update_buzzer();           //  B5. Quản lý buzzer\n'
'    }\n'
'\n'
'    uint16_t key = KeyScan();      // C. Đọc phím\n'
'    if (key) process_key(key);\n'
'}'
)

h2('4.2 A — __WDT_FEED_VALUE')
code('#define __WDT_FEED_VALUE  (SN_WDT->FEED = 0x5AFA5AFA)')
p('Macro ghi giá trị magic 0x5AFA5AFA vào thanh ghi FEED của WDT. '
  'Đây là "mật khẩu" phần cứng để reset bộ đếm WDT về 0. '
  'Nếu while(1) bị treo (bug, HardFault bị bỏ qua), '
  'WDT không được feed → timeout ~250ms → MCU tự reset. '
  'Đặt ở ĐẦU vòng lặp để bảo đảm main() đang chạy thực sự.')

h2('4.3 B1 — Digital_Scan(): quét LED 7 đoạn mỗi 1ms')
code(
'void Digital_Scan(void)\n'
'{\n'
'    SN_GPIO0->BCLR = 0xff;         // Tắt tất cả segment (tránh ghosting)\n'
'    SN_GPIO1->BCLR = 0xf << 9;    // Tắt tất cả COM\n'
'\n'
'    com_scan = (com_scan + 1) % 4; // Chuyển sang chữ số tiếp theo\n'
'\n'
'    switch(com_scan) {             // Bật COM tương ứng\n'
'        case 0: SN_GPIO1->BSET = 1 << 9;  break; // COM0 = P1.9\n'
'        case 1: SN_GPIO1->BSET = 1 << 10; break; // COM1 = P1.10\n'
'        case 2: SN_GPIO1->BSET = 1 << 11; break; // COM2 = P1.11\n'
'        case 3: SN_GPIO1->BSET = 1 << 12; break; // COM3 = P1.12\n'
'    }\n'
'    SN_GPIO0->BSET = segment_buff[com_scan]; // Bật segment từ buffer\n'
'}'
)
p('Mỗi 1ms, hàm này chuyển sang COM tiếp theo. '
  'Chu kỳ hoàn chỉnh 4 chữ số = 4ms → tần số quét = 250Hz. '
  'Thứ tự: tắt segment → tắt COM → chuyển COM → bật COM → bật segment. '
  'Thứ tự này tránh ghosting (bóng ma chữ số cũ).')
note('segment_buff[0]=COM0=chữ số hàng chục giờ, '
     '[1]=COM1=đơn vị giờ+DP, [2]=COM2=chục phút, [3]=COM3=đơn vị phút.')

h2('4.4 B2+B3 — Đếm giây qua g_ms_cnt')
code(
'g_ms_cnt++;\n'
'if (g_ms_cnt >= 1000U) {    // Đủ 1000ms = 1 giây\n'
'    g_ms_cnt = 0;\n'
'    tick_second();           // Xử lý tăng giây/phút/giờ\n'
'}'
)
p('g_ms_cnt là bộ đếm đơn giản 0–999. '
  'Không dùng timer thứ 3 vì CT16B1 đã cung cấp 1ms tick. '
  'Tích lũy 1000 tick = 1 giây thực tế.')

h2('4.5 B4 — Blink và Timeout')
code(
'if (g_state != STATE_NORMAL) {     // Chỉ hoạt động ở chế độ SET\n'
'    // BLINK: toggle mỗi 500ms\n'
'    g_blink_cnt++;\n'
'    if (g_blink_cnt >= BLINK_HALF_MS) {  // = 500\n'
'        g_blink_cnt = 0;\n'
'        g_blink_on = !g_blink_on;    // Đảo: 1→0 hoặc 0→1\n'
'        update_display();            // Áp dụng ngay vào segment_buff\n'
'    }\n'
'    // TIMEOUT: thoát về NORMAL sau 30 giây\n'
'    g_timeout_cnt++;\n'
'    if (g_timeout_cnt >= TIMEOUT_MS) {   // = 30000\n'
'        g_timeout_cnt = 0;\n'
'        g_blink_on = 1;\n'
'        g_state = STATE_NORMAL;\n'
'        start_pip();    // Pip báo đã timeout\n'
'        update_display();\n'
'    }\n'
'} else {\n'
'    g_blink_on = 1;   // NORMAL: luôn sáng hết\n'
'}'
)
note('g_blink_cnt và g_timeout_cnt chỉ tăng khi g_state != NORMAL. '
     'Khi nhấn phím, process_key() reset cả hai về 0 → gia hạn timeout, bắt đầu lại chu kỳ nháy.')

h2('4.6 C — KeyScan(): đọc phím với debounce')
code(
'// Trong while(1) — NGOÀI if(timer_1ms_flag)\n'
'uint16_t key = KeyScan();   // KeyScan tự debounce 50ms\n'
'if (key != 0) {\n'
'    process_key(key);        // Xử lý phím theo g_state\n'
'}'
)
p('KeyScan() được gọi trong MỖI vòng lặp (không phải mỗi 1ms). '
  'Hàm tự xử lý debounce 50ms nội bộ: phím phải giữ ổn định 50ms liên tục '
  'mới được coi là nhấn hợp lệ. Trả về 0 nếu không có phím.')

# ═══════════════════════════════════════════════════════════════════════════
h1('PHẦN 5 — main.c: tick_second() — ĐẾM THỜI GIAN')
# ═══════════════════════════════════════════════════════════════════════════
code(
'static void tick_second(void)\n'
'{\n'
'    if (g_state != STATE_NORMAL) return;  // ĐÓNG BĂNG khi đang cài đặt\n'
'\n'
'    g_sec++;\n'
'    if (g_sec >= 60) {\n'
'        g_sec = 0;\n'
'        g_min++;\n'
'        update_display();          // Cập nhật hiển thị mỗi khi phút thay đổi\n'
'\n'
'        if (g_min >= 60) {\n'
'            g_min = 0;\n'
'            g_hour++;\n'
'            if (g_hour >= 24) g_hour = 0;  // Wrap: 23:59 → 00:00\n'
'        }\n'
'\n'
'        // Kiểm tra báo thức CHỈ khi giây = 0 (tại HH:MM:00)\n'
'        if (g_hour == g_alarm_hour && g_min == g_alarm_min) {\n'
'            start_alarm_beep();\n'
'        }\n'
'    }\n'
'}'
)
p('Ba điểm quan trọng:')
li('Đồng hồ bị đóng băng (return sớm) khi đang ở SET_HOUR/SET_MIN/... '
   '→ người dùng cài giờ không bị đồng hồ nhảy số ở dưới')
li('update_display() chỉ gọi khi phút thay đổi (không phải mỗi giây) '
   '→ giảm tải không cần thiết')
li('Báo thức kích hoạt khi g_sec VỪA BẰNG 0 (tức là g_min vừa tăng lên giá trị khớp) '
   '→ không thể miss do race condition')

# ═══════════════════════════════════════════════════════════════════════════
h1('PHẦN 6 — main.c: process_key() — STATE MACHINE')
# ═══════════════════════════════════════════════════════════════════════════

h2('6.1 Phần chung — chạy trước khi xử lý state')
code(
'static void process_key(uint16_t key)\n'
'{\n'
'    start_pip();       // Pip 0.3s (bỏ qua nếu đang alarm)\n'
'    g_timeout_cnt = 0; // Reset timeout → người dùng còn tương tác\n'
'    g_blink_cnt   = 0; // Reset blink cycle → chữ số sáng ngay\n'
'    g_blink_on    = 1; // Bảo đảm chữ số sáng tức thì\n'
'    switch (g_state) { ... }\n'
'    update_display();  // Áp dụng thay đổi vào màn hình\n'
'}'
)
note('Lý do reset g_blink_cnt và g_blink_on=1: sau khi nhấn phím, '
     'người dùng cần thấy ngay giá trị mới (không bị tắt đúng lúc nhấn). '
     'Chữ số sẽ nháy tiếp sau 500ms.')

h2('6.2 Bảng chuyển trạng thái đầy đủ')
tbl(['Từ State','Phím nhấn','Đến State','Tác dụng phụ'],
[
 ['NORMAL',        'SW3 (SETUP)', 'SET_HOUR',       '(không)'],
 ['NORMAL',        'SW16 (ALARM)','SET_ALARM_HOUR', '(không)'],
 ['NORMAL',        'SW6/SW10',    'NORMAL',          'bị bỏ qua'],
 ['SET_HOUR',      'SW3',         'SET_MIN',         '(không)'],
 ['SET_HOUR',      'SW6',         'SET_HOUR',        'g_hour = (g_hour+1)%24'],
 ['SET_HOUR',      'SW10',        'SET_HOUR',        'g_hour = (g_hour==0)?23:g_hour-1'],
 ['SET_MIN',       'SW3',         'NORMAL',          'g_sec=0, g_ms_cnt=0'],
 ['SET_MIN',       'SW6',         'SET_MIN',         'g_min=(g_min+1)%60, g_sec=0'],
 ['SET_MIN',       'SW10',        'SET_MIN',         'g_min=(g_min==0)?59:g_min-1, g_sec=0'],
 ['SET_ALARM_HOUR','SW16',        'SET_ALARM_MIN',   '(không)'],
 ['SET_ALARM_HOUR','SW6',         'SET_ALARM_HOUR',  'g_alarm_hour=(g_alarm_hour+1)%24'],
 ['SET_ALARM_HOUR','SW10',        'SET_ALARM_HOUR',  'g_alarm_hour=(g_alarm_hour==0)?23:...'],
 ['SET_ALARM_MIN', 'SW16',        'NORMAL',          'save_alarm_to_eeprom()'],
 ['SET_ALARM_MIN', 'SW6',         'SET_ALARM_MIN',   'g_alarm_min=(g_alarm_min+1)%60'],
 ['SET_ALARM_MIN', 'SW10',        'SET_ALARM_MIN',   'g_alarm_min=(g_alarm_min==0)?59:...'],
 ['Bất kỳ SET','(30s không nhấn)','NORMAL',          'start_pip() — timeout'],
])

h2('6.3 Tại sao reset g_sec=0 và g_ms_cnt=0 khi xác nhận SET_MIN?')
p('Khi người dùng nhấn SW3 để xác nhận phút (chuyển về NORMAL), '
  'đồng hồ bắt đầu đếm từ HH:MM:00. '
  'Nếu không reset g_sec: ví dụ đang ở giây 55, '
  '5 giây sau đồng hồ tự động tăng phút thêm 1 lần nữa → sai 1 phút. '
  'g_ms_cnt=0 reset bộ đếm ms để giây đầu tiên đủ đúng 1000ms.')

# ═══════════════════════════════════════════════════════════════════════════
h1('PHẦN 7 — main.c: update_display() — CẬP NHẬT HIỂN THỊ')
# ═══════════════════════════════════════════════════════════════════════════
code(
'static void update_display(void)\n'
'{\n'
'    uint8_t h, m;\n'
'    uint8_t show_hh = 1, show_mm = 1;\n'
'\n'
'    // Bước 1: Chọn nguồn dữ liệu\n'
'    if (g_state == STATE_SET_ALARM_HOUR || g_state == STATE_SET_ALARM_MIN) {\n'
'        h = g_alarm_hour; m = g_alarm_min;  // Hiển thị giờ báo thức\n'
'    } else {\n'
'        h = g_hour; m = g_min;              // Hiển thị giờ hiện tại\n'
'    }\n'
'\n'
'    // Bước 2: Xác định nhóm nháy\n'
'    if (!g_blink_on) {\n'
'        if (g_state == STATE_SET_HOUR || g_state == STATE_SET_ALARM_HOUR)\n'
'            show_hh = 0;   // Tắt HH khi nháy ở chế độ chỉnh giờ\n'
'        else if (g_state == STATE_SET_MIN || g_state == STATE_SET_ALARM_MIN)\n'
'            show_mm = 0;   // Tắt MM khi nháy ở chế độ chỉnh phút\n'
'    }\n'
'\n'
'    // Bước 3: Ghi segment_buff\n'
'    segment_buff[0] = show_hh ? SEGMENT_TABLE[h/10]             : 0x00;\n'
'    segment_buff[1] = show_hh ? (SEGMENT_TABLE[h%10] | 0x80u)  : 0x00;\n'
'    segment_buff[2] = show_mm ? SEGMENT_TABLE[m/10]             : 0x00;\n'
'    segment_buff[3] = show_mm ? SEGMENT_TABLE[m%10]             : 0x00;\n'
'\n'
'    // Bước 4: LED D0\n'
'    if (g_state == STATE_SET_ALARM_HOUR || g_state == STATE_SET_ALARM_MIN)\n'
'        g_blink_on ? LED_ALARM_ON() : LED_ALARM_OFF();  // Nháy cùng chữ số\n'
'    else\n'
'        LED_ALARM_OFF();   // Tắt hẳn ở tất cả chế độ khác\n'
'}'
)
p('Điểm quan trọng trong segment_buff[1]: OR với 0x80u để bật bit 7 = chân P0.7 = đoạn DP. '
  'Đây chính là dấu chấm (.) giữa HH và MM trên LED 7 đoạn, tạo định dạng HH.MM. '
  'Khi show_hh=0 (đang nháy tắt), segment_buff[1]=0x00 → DP cũng tắt cùng.')

# ═══════════════════════════════════════════════════════════════════════════
h1('PHẦN 8 — main.c: QUẢN LÝ BUZZER')
# ═══════════════════════════════════════════════════════════════════════════

h2('8.1 buzzer_on() và buzzer_off()')
code(
'static void buzzer_on(void)\n'
'{\n'
'    SN_CT16B0->MR9 = BUZZER_PERIOD - 1;         // = 19999: chu kỳ PWM\n'
'    SN_CT16B0->MR0 = (BUZZER_PERIOD - 1) >> 1;  // = 9999: duty 50%\n'
'    SN_CT16B0->TMRCTRL = 0;  // Dừng timer, reset bộ đếm\n'
'    SN_CT16B0->TMRCTRL = 1;  // Khởi động lại\n'
'}\n'
'\n'
'static void buzzer_off(void)\n'
'{\n'
'    SN_CT16B0->MR0 = 0;  // Duty 0% → chân P3.0 luôn LOW → im lặng\n'
'}'
)
p('Tần số buzzer = HCLK / BUZZER_PERIOD = 12,000,000 / 20,000 = 600 Hz. '
  'MR9 là giá trị reload (chu kỳ đếm). MR0 là điểm đổi trạng thái PWM. '
  'Khi TC=0 → chân lên HIGH; khi TC=MR0 → chân xuống LOW; khi TC=MR9 → reset TC. '
  'Duty = MR0/MR9 = 50%. Tắt bằng cách đặt MR0=0 → chân luôn LOW.')

h2('8.2 update_buzzer() — logic pip và alarm mỗi 1ms')
code(
'static void update_buzzer(void)\n'
'{\n'
'    if (g_alarm_active) {\n'
'        // === Chế độ ALARM: pip-pip 5 giây ===\n'
'        if (g_alarm_remain_ms > 0) {\n'
'            g_alarm_remain_ms--;   // Đếm ngược từ 5000\n'
'            g_alarm_phase_ms--;    // Đếm ngược pha hiện tại từ 500\n'
'            if (g_alarm_phase_ms == 0) {\n'
'                g_alarm_phase_ms = ALARM_HALF_MS;    // Nạp lại 500ms\n'
'                g_alarm_phase_on = !g_alarm_phase_on; // Toggle pha\n'
'                g_alarm_phase_on ? buzzer_on() : buzzer_off();\n'
'            }\n'
'        } else {\n'
'            g_alarm_active = 0;\n'
'            buzzer_off();  // Kết thúc sau 5 giây\n'
'        }\n'
'    } else if (g_pip_ms > 0) {\n'
'        // === Chế độ PIP: đếm ngược 300ms ===\n'
'        g_pip_ms--;\n'
'        if (g_pip_ms == 0) buzzer_off();\n'
'    }\n'
'}'
)
note('Alarm có ưu tiên cao hơn pip: khi g_alarm_active=1, pip bị bỏ qua hoàn toàn. '
     'start_pip() kiểm tra if(!g_alarm_active) trước khi bật.')
tbl(['Trạng thái','g_alarm_active','g_pip_ms','Buzzer'],
[
 ['Im lặng',          '0', '0',    'OFF'],
 ['Pip ngắn (0.3s)',  '0', '300→0','ON trong 300ms rồi OFF'],
 ['Alarm (5s pip-pip)','1','0',    'ON 500ms → OFF 500ms → ... x5'],
])

# ═══════════════════════════════════════════════════════════════════════════
h1('PHẦN 9 — main.c: EEPROM READ/WRITE')
# ═══════════════════════════════════════════════════════════════════════════
code(
'static void load_alarm_from_eeprom(void)\n'
'{\n'
'    eeprom_read(EEPROM_READ_ADDR, EEPROM_AL_HOUR, &g_alarm_hour, 1);\n'
'    eeprom_read(EEPROM_READ_ADDR, EEPROM_AL_MIN,  &g_alarm_min,  1);\n'
'\n'
'    // EEPROM chưa ghi → trả về 0xFF (giá trị mặc định flash)\n'
'    if (g_alarm_hour >= 24) g_alarm_hour = 0;\n'
'    if (g_alarm_min  >= 60) g_alarm_min  = 0;\n'
'}\n'
'\n'
'static void save_alarm_to_eeprom(void)\n'
'{\n'
'    eeprom_write(EEPROM_WRITE_ADDR, EEPROM_AL_HOUR, &g_alarm_hour, 1);\n'
'    eeprom_write(EEPROM_WRITE_ADDR, EEPROM_AL_MIN,  &g_alarm_min,  1);\n'
'}'
)
p('EEPROM_READ_ADDR = 0xA1: byte địa chỉ I2C với bit R/W=1. '
  'EEPROM_WRITE_ADDR = 0xA0: byte địa chỉ I2C với bit R/W=0. '
  'Lệnh đọc 1 byte từ ô 0x00: Start → 0xA0 → 0x00 → RepeatedStart → 0xA1 → data → NACK → Stop. '
  'Kiểm tra 0xFF: AT24C02 mới xuất xưởng, tất cả ô nhớ = 0xFF → '
  'giờ báo thức 0xFF:0xFF = 255:255 → không hợp lệ → đặt về 00:00.')

# ═══════════════════════════════════════════════════════════════════════════
h1('PHẦN 10 — CT16B1.c: TIMER 1ms')
# ═══════════════════════════════════════════════════════════════════════════
code(
'void CT16B1_Init(void)\n'
'{\n'
'    __CT16B1_ENABLE;                  // Bật clock APB cho CT16B1\n'
'\n'
'    SN_CT16B1->MR9 = 12*1000 - 1;   // MR9 = 11999\n'
'    // Tính: HCLK=12MHz, không có prescaler → 1 tick = 1/12MHz = 83.3ns\n'
'    // 12000 ticks × 83.3ns = 1,000,000 ns = 1ms\n'
'\n'
'    SN_CT16B1->MCTRL = (1<<30)|(1<<29); // bit30: reset TC khi match MR9\n'
'    //                                   // bit29: bật IRQ khi match MR9\n'
'\n'
'    SN_CT16B1->TMRCTRL = (1<<1);     // Reset Timer Counter về 0\n'
'    while(SN_CT16B1->TMRCTRL & (1<<1)); // Chờ reset xong\n'
'    SN_CT16B1->TMRCTRL = 1;          // Start timer\n'
'\n'
'    NVIC_EnableIRQ(CT16B1_IRQn);     // Bật ngắt CT16B1 trong NVIC\n'
'}\n'
'\n'
'void CT16B1_IRQHandler(void)\n'
'{\n'
'    uint16_t ris = SN_CT16B1->RIS;   // Đọc cờ ngắt\n'
'    if (ris & (1<<5))                 // bit5 = MR9 interrupt flag\n'
'        timer_1ms_flag = 1;           // Báo cho main loop\n'
'    SN_CT16B1->IC = ris;              // PHẢI xóa cờ, không thì ISR bị gọi liên tục\n'
'}'
)
tbl(['Thanh ghi','Giá trị','Ý nghĩa'],
[
 ['MR9',    '11999',       'Match value: TC sẽ reset khi đạt 11999'],
 ['MCTRL bit30','1',       'Reset TC về 0 khi TC==MR9 (chế độ auto-reload)'],
 ['MCTRL bit29','1',       'Sinh ngắt khi TC==MR9'],
 ['TMRCTRL bit0','1',      'Bật timer chạy'],
 ['RIS bit5','set bởi HW', 'Cờ match MR9 đang pending'],
 ['IC',     'ghi = ris',   'Clear cờ (ghi 1 vào bit để xóa)'],
])

# ═══════════════════════════════════════════════════════════════════════════
h1('PHẦN 11 — CT16B0.c: BUZZER PWM')
# ═══════════════════════════════════════════════════════════════════════════
code(
'void CT16B0_Init(void)\n'
'{\n'
'    __CT16B0_ENABLE;                  // Bật clock APB cho CT16B0\n'
'    SN_CT16B0->MR9  = 12*1000 - 1;  // Giá trị khởi tạo (sẽ bị ghi đè bởi buzzer_on)\n'
'    SN_CT16B0->MR0  = 0;             // Duty = 0% → tắt\n'
'    SN_CT16B0->MCTRL = (1<<30);      // Reset TC khi match MR9\n'
'    SN_CT16B0->PWMCTRL = (1<<0)      // Bật PWM cho MR0\n'
'                        |(1<<4)      // PWM mode 2 (active HIGH)\n'
'                        |(1<<20);    // Cho phép output ra chân vật lý\n'
'    SN_CT16B0->TMRCTRL = (1<<1);    // Reset TC\n'
'    while(SN_CT16B0->TMRCTRL & (1<<1));\n'
'    SN_CT16B0->TMRCTRL = 1;         // Start timer\n'
'}\n'
'\n'
'// Sau init: gán chân PWM0 ra P3.0 qua PFPA\n'
'SN_PFPA->CT16B0_b.PWM0 = 1;   // PWM0 → P3.0 (option 1)'
)
note('CT16B0_Init() đặt MR9=11999 nhưng buzzer_on() sẽ ghi đè MR9=19999 (600Hz). '
     'Init chỉ cần PWMCTRL và TMRCTRL chạy đúng. '
     'bit20 (IO output enable) phải được set mới có tín hiệu ra chân.')

# ═══════════════════════════════════════════════════════════════════════════
h1('PHẦN 12 — GPIO.c: CẤU HÌNH CHÂN')
# ═══════════════════════════════════════════════════════════════════════════
code(
'void GPIO_Init(void)\n'
'{\n'
'    // --- Phím: INPUT ---\n'
'    SN_GPIO1->MODE = 0xf << 4;   // P1.4~P1.7: input (KEY ROW) — MODE=00\n'
'    SN_GPIO2->MODE = 0xf << 4;   // P2.4~P2.7: input (KEY COL) — MODE=00\n'
'\n'
'    // --- LED 7 đoạn SEGMENT: OUTPUT ---\n'
'    SN_GPIO0->MODE = 0xff;       // P0.0~P0.7: output push-pull — MODE=01\n'
'    SN_GPIO0->BCLR = 0xff;       // Tắt tất cả segment ban đầu\n'
'\n'
'    // --- LED 7 đoạn COM: OUTPUT ---\n'
'    SN_GPIO1->MODE |= (0x0f << 9);  // P1.9~P1.12: output — COM0~COM3\n'
'    SN_GPIO1->BCLR  = (0x0f << 9);  // Tắt tất cả COM ban đầu\n'
'\n'
'    // --- LED D0 (báo thức): OUTPUT ---\n'
'    SN_GPIO3->MODE_b.MODE8 = 1;  // P3.8: output\n'
'    SET_LED0_OFF;                 // = SN_GPIO3->BSET=(1<<8) — tắt (active-LOW)\n'
'\n'
'    // --- LED D1 (dự phòng): OUTPUT ---\n'
'    SN_GPIO3->MODE_b.MODE9 = 1;\n'
'    SET_LED1_OFF;\n'
'}'
)
tbl(['Chân GPIO','Chức năng','Hướng','Ghi chú'],
[
 ['P0.0–P0.7','SEG a–g, DP','Output','Dữ liệu 7 đoạn (bit0=a, bit7=DP)'],
 ['P1.9–P1.12','COM0–COM3','Output','Chọn chữ số đang sáng'],
 ['P1.4–P1.7','KEY ROW 1–4','Input','Hàng ma trận phím, pull-up nội'],
 ['P2.4–P2.7','KEY COL 1–4','Input','Cột ma trận phím, pull-up nội'],
 ['P3.0','CT16B0 PWM0','Output (PFPA)','Buzzer passive ~600Hz'],
 ['P3.8','LED D0 (alarm)','Output','Active-LOW: BCLR=sáng, BSET=tắt'],
 ['P0.10','I2C0 SCL','Output (PFPA)','Clock I2C, open-drain'],
 ['P0.11','I2C0 SDA','I/O (PFPA)','Data I2C, open-drain'],
])

# ═══════════════════════════════════════════════════════════════════════════
h1('PHẦN 13 — Segment.c: BẢNG MÃ VÀ QUÉT LED')
# ═══════════════════════════════════════════════════════════════════════════

h2('13.1 SEGMENT_TABLE[]')
code(
'// Định nghĩa đoạn:\n'
'#define SEG_A 0x01  // bit0 → P0.0\n'
'#define SEG_B 0x02  // bit1 → P0.1\n'
'#define SEG_C 0x04  // bit2 → P0.2\n'
'#define SEG_D 0x08  // bit3 → P0.3\n'
'#define SEG_E 0x10  // bit4 → P0.4\n'
'#define SEG_F 0x20  // bit5 → P0.5\n'
'#define SEG_G 0x40  // bit6 → P0.6\n'
'#define SEG_H 0x80  // bit7 → P0.7 (DP - Decimal Point)\n'
'\n'
'const uint8_t SEGMENT_TABLE[] = {\n'
'    (SEG_A|SEG_B|SEG_C|SEG_D|SEG_E|SEG_F),          // 0 = 0x3F = 0b00111111\n'
'    (SEG_B|SEG_C),                                   // 1 = 0x06 = 0b00000110\n'
'    (SEG_A|SEG_B|SEG_D|SEG_E|SEG_G),                // 2 = 0x5B\n'
'    (SEG_A|SEG_B|SEG_C|SEG_D|SEG_G),                // 3 = 0x4F\n'
'    (SEG_B|SEG_C|SEG_F|SEG_G),                      // 4 = 0x66\n'
'    (SEG_A|SEG_C|SEG_D|SEG_F|SEG_G),                // 5 = 0x6D\n'
'    (SEG_A|SEG_C|SEG_D|SEG_E|SEG_F|SEG_G),          // 6 = 0x7D\n'
'    (SEG_A|SEG_B|SEG_C),                             // 7 = 0x07\n'
'    (SEG_A|SEG_B|SEG_C|SEG_D|SEG_E|SEG_F|SEG_G),   // 8 = 0x7F\n'
'    (SEG_A|SEG_B|SEG_C|SEG_D|SEG_F|SEG_G),          // 9 = 0x6F\n'
'};'
)
note('Cách đọc: số 2 cần đoạn a(trên), b(phải trên), d(dưới), e(trái dưới), g(giữa). '
     'Tương ứng: SEG_A|SEG_B|SEG_D|SEG_E|SEG_G = 0x01|0x02|0x08|0x10|0x40 = 0x5B. '
     'Ghi SN_GPIO0->BSET=0x5B: bật P0.0, P0.1, P0.3, P0.4, P0.6 → chữ số 2 sáng.')

h2('13.2 Digital_Scan() — logic quét')
code(
'void Digital_Scan(void)\n'
'{\n'
'    // Bước 1: Tắt tất cả để tránh ghosting\n'
'    SN_GPIO0->BCLR = 0xff;       // Tắt segment\n'
'    SN_GPIO1->BCLR = 0xf << 9;  // Tắt COM\n'
'\n'
'    // Bước 2: Chuyển sang chữ số tiếp theo (0→1→2→3→0)\n'
'    com_scan = (++com_scan >= 4) ? 0 : com_scan;\n'
'\n'
'    // Bước 3: Bật COM tương ứng\n'
'    SN_GPIO1->BSET = (1 << (9 + com_scan));  // COM0=P1.9, COM1=P1.10...\n'
'\n'
'    // Bước 4: Xuất mã 7 đoạn từ buffer\n'
'    SN_GPIO0->BSET = segment_buff[com_scan];\n'
'}'
)
note('Thứ tự PHẢI là: tắt SEG → tắt COM → đổi COM → bật COM → bật SEG. '
     'Nếu đổi thứ tự: bật COM trước khi tắt SEG → LED com_scan cũ nhận data của com_scan mới → ghosting.')

# ═══════════════════════════════════════════════════════════════════════════
h1('PHẦN 14 — EEPROM.c: GIAO TIẾP I2C')
# ═══════════════════════════════════════════════════════════════════════════

h2('14.1 eeprom_write() — ghi 1 byte')
code(
'// Giao thức: [START][0xA0 W][ACK][reg addr][ACK][data][ACK][STOP]\n'
'void eeprom_write(uint8_t addr, uint8_t reg, uint8_t *dat, uint16_t length)\n'
'{\n'
'    I2C0_Start();\n'
'    if(I2C_write_byte(addr)==I2C_NACK_FALG){ I2C0_Stop(); return; }  // Gửi 0xA0\n'
'    if(I2C_write_byte(reg)==I2C_NACK_FALG) { I2C0_Stop(); return; }  // Gửi địa chỉ ô\n'
'    while(length--) {\n'
'        if(I2C_write_byte(*dat++)==I2C_NACK_FALG){ I2C0_Stop(); return; }\n'
'    }\n'
'    I2C0_Stop();\n'
'}'
)

h2('14.2 eeprom_read() — đọc 1 byte (Repeated START)')
code(
'// Giao thức: [START][0xA0 W][ACK][reg][ACK][rSTART][0xA1 R][ACK][data][NACK][STOP]\n'
'void eeprom_read(uint8_t addr, uint8_t reg, uint8_t *dat, uint16_t length)\n'
'{\n'
'    I2C0_Start();\n'
'    I2C_write_byte(addr & 0xFE);  // 0xA0: Write mode — set địa chỉ ô nhớ\n'
'    I2C_write_byte(reg);          // Ô nhớ cần đọc (0x00 hoặc 0x01)\n'
'    I2C0_Start();                 // Repeated START — không có STOP ở giữa\n'
'    I2C_write_byte(addr);         // 0xA1: Read mode\n'
'    while(length > 1)\n'
'        *dat++ = I2C_read_byte(I2C_ACK_FALG);   // ACK mỗi byte trừ byte cuối\n'
'    *dat++ = I2C_read_byte(I2C_NACK_FALG);      // NACK byte cuối → báo kết thúc\n'
'    I2C0_Stop();\n'
'}'
)
p('Tại sao phải Repeated START? EEPROM AT24C02 cần "Random Read": '
  '(1) Ghi địa chỉ ô cần đọc (dùng write frame). '
  '(2) Repeated START để không nhả bus. '
  '(3) Gửi địa chỉ thiết bị với R/W=1. '
  '(4) Nhận data. Nếu dùng STOP sau bước 1, EEPROM quên địa chỉ ô, '
  'đọc lại từ đầu (Sequential Read từ địa chỉ 0 hoặc địa chỉ sau cùng).')

# ═══════════════════════════════════════════════════════════════════════════
h1('PHẦN 15 — SysTick_1ms.c: TIMER MÔ PHỎNG')
# ═══════════════════════════════════════════════════════════════════════════
code(
'#define SYSTICK_HCLK_HZ  12000000UL\n'
'#define SYSTICK_FREQ_HZ      1000UL\n'
'#define SYSTICK_RELOAD   ((SYSTICK_HCLK_HZ/SYSTICK_FREQ_HZ)-1UL)  // = 11999\n'
'\n'
'volatile uint8_t timer_1ms_flag = 0;  // volatile: ghi từ ISR, đọc từ main()\n'
'\n'
'void SysTick_Handler(void)\n'
'{\n'
'    timer_1ms_flag = 1;  // Không làm gì thêm — nhẹ nhàng nhất có thể\n'
'}\n'
'\n'
'void SysTick_1ms_Init(void)\n'
'{\n'
'    SysTick->CTRL = 0;               // Tắt trước khi cấu hình\n'
'    SysTick->LOAD = SYSTICK_RELOAD;  // = 11999\n'
'    SysTick->VAL  = 0;               // Reset bộ đếm hiện tại về 0\n'
'    SysTick->CTRL = 0x07;            // ENABLE=1, TICKINT=1, CLKSOURCE=1(HCLK)\n'
'}'
)
p('SysTick là bộ đếm 24-bit ĐẾM NGƯỢC từ LOAD về 0. '
  'Mỗi khi về 0: tự nạp lại LOAD, set COUNTFLAG, gọi SysTick_Handler(). '
  'CTRL bit2=1: dùng HCLK (không chia đôi). '
  'CTRL bit1=1: gọi ISR khi về 0. '
  'CTRL bit0=1: bật SysTick. '
  'Dùng trong bản simulator vì Keil mô phỏng SysTick (ARM core chuẩn), '
  'trong khi CT16B1 (ngoại vi SN32F407) không được mô phỏng.')

# ═══════════════════════════════════════════════════════════════════════════
h1('PHẦN 16 — LUỒNG DỮ LIỆU ĐẦY ĐỦ: TỪ PHÍM → LED')
# ═══════════════════════════════════════════════════════════════════════════
p('Ví dụ: người dùng nhấn SW3 khi đang NORMAL → đồng hồ vào SET_HOUR, '
  'chữ số giờ bắt đầu nháy.')
tbl(['Bước','Xảy ra ở đâu','Chi tiết'],
[
 ['1','Phần cứng','SW3 (P2.4/col) được nhấn → nối đất chân ROW tương ứng'],
 ['2','KeyScan.c','Digital_Scan() phát hiện ROW×COL → debounce 50ms → trả KEY_3=0x0014'],
 ['3','main() while(1)','key=KeyScan()=0x0014 ≠ 0 → gọi process_key(0x0014)'],
 ['4','process_key()','start_pip(): g_pip_ms=300, buzzer_on()'],
 ['5','process_key()','g_timeout_cnt=0, g_blink_cnt=0, g_blink_on=1'],
 ['6','process_key() switch','case STATE_NORMAL: key==KEY_SETUP → g_state=STATE_SET_HOUR(=1)'],
 ['7','process_key()','update_display(): show_hh=1 show_mm=1, ghi segment_buff[]'],
 ['8','CT16B1 ISR (sau 1ms)','timer_1ms_flag=1'],
 ['9','main() 1ms block','Digital_Scan(): tắt SEG → đổi COM → bật COM → bật segment_buff[com_scan]'],
 ['10','main() 1ms block','g_blink_cnt++ (=1)'],
 ['11','(500ms sau)','g_blink_cnt=500 → g_blink_on=0 → update_display(): show_hh=0 → segment_buff[0]=0, [1]=0'],
 ['12','Digital_Scan()','P0.0-P0.7 = 0x00 → tất cả segment COM0,COM1 tắt → chữ số giờ tắt'],
 ['13','(500ms sau nữa)','g_blink_on=1 → segment_buff khôi phục → chữ số giờ sáng lại'],
])

blank()
p('Vòng lặp cứ thế cho đến khi: nhấn phím tiếp (reset g_blink_cnt) '
  'hoặc timeout 30s (g_state về NORMAL).')

out = r'C:\Sonix\Clock_SN32F407\Tài_liệu_giải_thích_code.docx'
doc.save(out)
print('Done:', out)
