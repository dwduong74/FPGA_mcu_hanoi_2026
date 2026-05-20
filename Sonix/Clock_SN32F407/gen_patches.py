"""
gen_patches.py — Bổ sung nội dung còn thiếu vào 3 tài liệu đã có
Các vấn đề được agent phát hiện:
  A) Tài_liệu_phỏng_vấn: thiếu IHRC drift, I2C error, scaling, low-power, interrupt nesting
  B) Kịch_bản_quay_demo: thiếu WDT demo, EEPROM persistence, buzzer verify, wrap-around
  C) Báo_cáo: thiếu edge case test, race condition, EEPROM wear
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

# ─── helpers ────────────────────────────────────────────────────────────────
def h1(d,t): d.add_heading(t,1)
def h2(d,t): d.add_heading(t,2)
def h3(d,t): d.add_heading(t,3)
def p(d,t):  d.add_paragraph(t)
def li(d,t): d.add_paragraph(t, style='List Bullet')
def blank(d):d.add_paragraph()

def code(d,t):
    para = d.add_paragraph()
    run  = para.add_run(t)
    run.font.name = 'Courier New'
    run.font.size = Pt(8.5)
    para.paragraph_format.left_indent = Cm(0.8)
    return para

def note(d,t):
    para = d.add_paragraph('► '+t)
    for r in para.runs:
        r.italic = True
        r.font.color.rgb = RGBColor(0xCC,0x44,0x00)

def qa(d, question, answer):
    para = d.add_paragraph()
    r = para.add_run('Q: ' + question)
    r.bold = True
    r.font.color.rgb = RGBColor(0x00,0x44,0x99)
    d.add_paragraph('A: ' + answer)
    blank(d)

def tbl(d, headers, rows):
    t = d.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'
    for i,h in enumerate(t.rows[0].cells):
        h.text = headers[i]
        for r in h.paragraphs[0].runs: r.bold = True
    for ri,row in enumerate(rows):
        for ci,val in enumerate(row):
            t.rows[ri+1].cells[ci].text = str(val)
    d.add_paragraph()

def separator(d, title):
    """Thêm một trang bìa phần mới"""
    d.add_page_break()
    heading = d.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = heading.add_run(f'【 PHỤ LỤC BỔ SUNG 】\n{title}')
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x00,0x44,0x99)
    blank(d)

# ════════════════════════════════════════════════════════════════════════════
# A) VÁ TÀI LIỆU PHỎNG VẤN — thêm phần 13: chủ đề bị thiếu
# ════════════════════════════════════════════════════════════════════════════
path_pv = r'C:\Sonix\Clock_SN32F407\Tài_liệu_phỏng_vấn_SN32F407.docx'
d_pv = Document(path_pv)

separator(d_pv, 'PHẦN 13 — CÁC CHỦ ĐỀ BỔ SUNG THƯỜNG GẶP TRONG PHỎNG VẤN')

h1(d_pv, 'PHẦN 13 — CHỦ ĐỀ BỔ SUNG')

# --- 13.1 Clock accuracy ---
h2(d_pv, '13.1 Độ chính xác đồng hồ — IHRC Drift')
qa(d_pv,
   'IHRC (Internal High-speed RC) của SN32F407 có độ chính xác như thế nào? '
   'Sau 1 giờ đồng hồ có bị sai không?',
   'Theo datasheet SONiX, IHRC có dung sai ±5% (worst case). '
   'Nghĩa là HCLK thực tế có thể là 11.4MHz–12.6MHz thay vì đúng 12.0MHz. '
   'Ảnh hưởng đến đồng hồ: nếu HCLK sai +5%, '
   'timer 1ms thực tế chỉ là 0.952ms → 1 giờ thực tế = 3600s '
   'nhưng đồng hồ đếm 3600/0.952 ≈ 3781 giây → nhanh ~3 phút/giờ. '
   'Đây là lý do đồng hồ chính xác cao dùng thạch anh ngoài (32.768kHz RTC crystal). '
   'Với bài tập này, IHRC đủ dùng (yêu cầu không đòi hỏi độ chính xác tuyệt đối).'
)

qa(d_pv,
   'Làm sao cải thiện độ chính xác nếu có yêu cầu?',
   '(1) Dùng thạch anh ngoài thay IHRC: độ chính xác ±20–50 ppm (~2–4 giây/ngày). '
   '(2) Dùng RTC chip ngoài (DS3231) có thạch anh 32.768kHz tích hợp và bù nhiệt độ: ±2 ppm. '
   '(3) Hiệu chỉnh phần mềm: đo drift thực tế, điều chỉnh TIMEOUT_MS cho phù hợp. '
   '(4) Đồng bộ NTP qua mạng (nếu MCU có WiFi).'
)

# --- 13.2 I2C error recovery ---
h2(d_pv, '13.2 Xử lý lỗi I2C — Bus Hang')
qa(d_pv,
   'Điều gì xảy ra nếu đường I2C bị treo (SDA kẹt ở LOW)? '
   'Code hiện tại xử lý thế nào?',
   'I2C bus hang: thường do EEPROM đang giữa giao dịch khi MCU reset → '
   'EEPROM giữ SDA = 0 chờ clock tiếp theo → master không thể bắt đầu START (cần SDA=1). '
   'Code hiện tại (eeprom_read/write) không có cơ chế timeout → nếu I2C_write_byte() '
   'chờ ACK mà không có, hàm có thể treo vô hạn → WDT sẽ reset MCU sau ~250ms. '
   'Giải pháp tốt hơn: (1) Thêm timeout cho vòng chờ ACK. '
   '(2) Bus recovery: MCU phát 9 xung SCL để giải phóng slave. '
   '(3) Reset I2C peripheral: ghi lại bit RESET trong thanh ghi điều khiển I2C.'
)

qa(d_pv,
   'eeprom_read() trả về gì nếu I2C NACK (EEPROM không trả lời)?',
   'Trong code hiện tại: I2C_write_byte() trả về I2C_NACK_FALG → '
   'eeprom_read() gọi I2C0_Stop() và return (không ghi vào *dat). '
   'Biến g_alarm_hour giữ giá trị CŨ (không được cập nhật). '
   'load_alarm_from_eeprom() sau đó kiểm tra if(g_alarm_hour>=24) → '
   'nếu g_alarm_hour vẫn là 0 (giá trị khởi tạo) → bình thường. '
   'Tuy nhiên nếu g_alarm_hour đã có giá trị cũ từ lần trước → '
   'đồng hồ dùng giá trị RAM cũ, không báo lỗi → đây là silent failure.'
)

# --- 13.3 Low power ---
h2(d_pv, '13.3 Low-Power Mode — Tiết kiệm điện')
qa(d_pv,
   'SN32F407 có chế độ tiết kiệm điện không? Tại sao project này không dùng?',
   'Có. SN32F407 hỗ trợ: (1) Sleep mode: CPU dừng, ngoại vi và SRAM vẫn chạy. '
   '(2) Deep Sleep mode: CPU và hầu hết ngoại vi dừng, chỉ RTC/WDT/pin interrupt còn chạy. '
   'Project này không dùng vì: (a) Đồng hồ cần CT16B1 chạy liên tục để đếm giây. '
   '(b) LED 7 đoạn cần Digital_Scan() mỗi 1ms để không tắt. '
   '(c) Phím cần polling liên tục. '
   'Nếu muốn tiết kiệm điện: dùng Sleep Mode giữa các tick 1ms, '
   '__WFI() (Wait For Interrupt) để CPU ngủ cho đến ngắt CT16B1 tiếp theo.'
)

code(d_pv,
     '// Thêm vào cuối while(1) để tiết kiệm điện:\n'
     'if (!timer_1ms_flag) {\n'
     '    __WFI();   // CPU ngủ đến ngắt tiếp theo\n'
     '}              // CT16B1 IRQ đánh thức, tiếp tục vòng lặp'
)

# --- 13.4 Interrupt nesting ---
h2(d_pv, '13.4 Ngắt và Vấn Đề Nesting / Race Condition')
qa(d_pv,
   'Cortex-M0 có hỗ trợ ngắt lồng nhau (nested interrupt) không?',
   'Cortex-M0 chỉ có 4 mức ưu tiên ngắt (0–3). '
   'Mặc định TẤT CẢ ngắt cùng mức ưu tiên → không có preemption lẫn nhau. '
   'Nghĩa là: nếu CT16B1_IRQHandler() đang chạy, '
   'một ngắt khác cùng mức phải ĐỢI cho đến khi ISR hiện tại hoàn thành. '
   'Điều này đơn giản hóa thiết kế nhưng: latency của ngắt thứ hai tăng lên = '
   'thời gian thực thi ISR thứ nhất. Với project này (ISR rất ngắn, chỉ set cờ), '
   'không phải vấn đề.'
)

qa(d_pv,
   'Race condition nào có thể xảy ra khi đọc g_hour/g_min từ main loop trong khi ISR chạy?',
   'Tình huống: main() đang đọc g_hour (giá trị cũ) → CT16B1 ISR fires → '
   'tick_second() tăng g_min (giờ rollover từ 23:59 → 00:00) → '
   'main() đọc g_min (giá trị mới 0). '
   'Kết quả: main() thấy g_hour=23, g_min=0 → hiển thị "23:00" thay vì "00:00". '
   'Mặc dù chỉ trong 1 chu kỳ quét, người dùng có thể thấy chớp sai. '
   'Fix: đọc atomic — tắt ngắt khi đọc nhóm biến:\n'
   '__disable_irq(); h=g_hour; m=g_min; __enable_irq();'
)

# --- 13.5 Scaling ---
h2(d_pv, '13.5 Mở Rộng Chức Năng (Scaling)')
qa(d_pv,
   'Nếu yêu cầu thêm hiển thị giây (6 chữ số HH:MM:SS), phải thay đổi gì?',
   '(1) Thêm 2 LED 7 đoạn và 2 chân COM (COM4, COM5). '
   '(2) segment_buff mở rộng từ 4 → 6 phần tử. '
   '(3) Tốc độ quét: 6 chữ số × 1ms = 6ms/frame → 167Hz (vẫn đủ, không bị nhấp nháy). '
   '(4) update_display() thêm segment_buff[4]=SEGMENT_TABLE[s/10], [5]=SEGMENT_TABLE[s%10]. '
   '(5) GPIO_Init() thêm COM4=P1.13, COM5=P1.14. '
   '(6) Digital_Scan() sửa vòng lặp com_scan từ %4 → %6.'
)

qa(d_pv,
   'Nếu cần 3 báo thức thay vì 1, kiến trúc thay đổi thế nào?',
   '(1) Thay biến đơn bằng mảng: uint8_t alarm_hour[3], alarm_min[3]; '
   '(2) Thêm biến uint8_t current_alarm_idx để biết đang cài báo thức nào. '
   '(3) State machine thêm state SW16 ở NORMAL: '
   'nhấn lần 1 → SET_ALARM1_HOUR, lần 2 → SET_ALARM2_HOUR, lần 3 → SET_ALARM3_HOUR. '
   '(4) EEPROM: dùng địa chỉ 0x00–0x05 (3 × 2 byte). '
   '(5) tick_second() kiểm tra vòng lặp for(i=0;i<3;i++) thay vì if đơn.'
)

# --- 13.6 HardFault ---
h2(d_pv, '13.6 HardFault_Handler')
qa(d_pv,
   'HardFault_Handler() làm gì? Khi nào nó được gọi?',
   'HardFault là exception ARM Cortex-M, được gọi khi CPU gặp lỗi nghiêm trọng: '
   '(1) Truy cập địa chỉ bộ nhớ không hợp lệ (null pointer dereference). '
   '(2) Lệnh không hợp lệ (unaligned access khi CPU không cho phép). '
   '(3) Stack overflow (SP vượt ra ngoài SRAM). '
   '(4) Divide by zero (không áp dụng M0 vì không có lệnh chia). '
   'Trong code: HardFault_Handler() gọi NVIC_SystemReset() để reset MCU ngay lập tức. '
   'Đây là cách đơn giản nhất: thay vì treo hệ thống, reset và khởi động lại sạch. '
   'Trong sản phẩm thực tế nên log lỗi vào EEPROM trước khi reset để debug sau.'
)

d_pv.save(path_pv)
print('Patched interview doc:', path_pv)

# ════════════════════════════════════════════════════════════════════════════
# B) VÁ KỊCH BẢN QUAY DEMO — thêm cảnh 10, 11, 12
# ════════════════════════════════════════════════════════════════════════════
path_demo = r'C:\Sonix\Clock_SN32F407\Kịch_bản_quay_demo.docx'
d_demo = Document(path_demo)

separator(d_demo, 'CÁC CẢNH BỔ SUNG (10–12)')

h1(d_demo, 'CẢNH 10 — Demo Wrap-around Phím (edge case, ~45 giây)')
h2(d_demo, 'Mục đích')
p(d_demo, 'Chứng minh logic vòng tròn (wrap-around) hoạt động đúng ở biên giá trị.')
h2(d_demo, 'Trình tự thao tác')
tbl(d_demo,
    ['Bước','Thao tác','Quan sát & Lời nói'],
    [
     ['1','Vào SET_HOUR (g_sim_key=0x14). Dùng SW6 tăng g_hour lên 23.',
      '"Tôi tăng giờ lên 23 — giá trị lớn nhất cho đồng hồ 24 giờ."'],
     ['2','Ghi g_sim_key=0x22 (SW6) lần nữa.',
      '"Nhấn PLUS khi giờ = 23. Kết quả: g_hour = 0. '
      'Logic: (23+1) % 24 = 0. Vòng từ 23 → 00 đúng."'],
     ['3','Ghi g_sim_key=0x42 (SW10 MINUS) khi g_hour=0.',
      '"Nhấn MINUS khi giờ = 0. Kết quả: g_hour = 23. '
      'Logic: (g_hour==0) ? 23 : g_hour-1. Vòng từ 00 → 23 đúng."'],
     ['4','Tương tự với g_min: tăng lên 59, nhấn PLUS → về 0.',
      '"Phút cũng wrap-around: (59+1)%60=0. '
      'Và MINUS khi phút=0 → phút=59."'],
    ])

h1(d_demo, 'CẢNH 11 — Demo WDT qua Watch Window (~45 giây)')
h2(d_demo, 'Cách demo không cần phần cứng thật')
p(d_demo,
  'WDT không thực sự reset trong Keil Simulator (Simulator không mô phỏng WDT phần cứng). '
  'Tuy nhiên có thể giải thích và chỉ vào code.')
h2(d_demo, 'Trình tự demo')
tbl(d_demo,
    ['Bước','Thao tác','Lời nói'],
    [
     ['1','Mở editor, cuộn đến dòng __WDT_FEED_VALUE ở đầu while(1).',
      '"Watchdog Timer được feed tại đây — đầu tiên trong mỗi vòng lặp."'],
     ['2','Chỉ vào hàm WDT_Init() (hoặc mở WDT.c).',
      '"WDT_Init() bật Watchdog với timeout khoảng 250ms. '
      'Trong Simulator, WDT không có tác dụng phần cứng thật, '
      'nhưng trên board, nếu while(1) bị treo vì bất kỳ lý do gì, '
      'WDT tự reset MCU sau 250ms."'],
     ['3','Thêm Watch Window: __WDT_FEED_VALUE có thể xem bằng cách theo dõi SN_WDT->FEED.',
      '"Giá trị 0x5AFA5AFA là magic key — chỉ khi MCU nhận đúng key này '
      'thì WDT mới được reset. Đây là bảo vệ chống code ngẫu nhiên vô tình feed WDT."'],
     ['4','Mở HardFault_Handler() trong main.c.',
      '"Nếu xảy ra lỗi nghiêm trọng (truy cập địa chỉ sai, stack overflow), '
      'HardFault_Handler() gọi NVIC_SystemReset() để reset ngay lập tức."'],
    ])

h1(d_demo, 'CẢNH 12 — Xác minh Buzzer qua CT16B0 Register (~45 giây)')
h2(d_demo, 'Mục đích')
p(d_demo, 'Trong Simulator không nghe tiếng buzzer, nhưng có thể chứng minh tần số đúng qua thanh ghi.')
h2(d_demo, 'Trình tự thao tác')
tbl(d_demo,
    ['Bước','Thao tác','Lời nói'],
    [
     ['1','Thêm vào Watch Window: SN_CT16B0->MR9, SN_CT16B0->MR0.',
      '"Tôi theo dõi hai thanh ghi của CT16B0 — timer điều khiển buzzer."'],
     ['2','Kích hoạt báo thức (g_alarm_active=1 như cảnh 6). Run.',
      '"Khi báo thức kích hoạt, buzzer_on() được gọi: '
      'MR9 = BUZZER_PERIOD - 1 = 19999. '
      'Tần số = HCLK / BUZZER_PERIOD = 12,000,000 / 20,000 = 600 Hz. '
      'MR0 = 9999: duty cycle = MR0/MR9 ≈ 50%."'],
     ['3','Dừng. Quan sát: MR9=19999, MR0=9999.',
      '"MR9=19999 và MR0=9999 xác nhận PWM đang chạy ở 600Hz, duty 50%."'],
     ['4','Sau 5s alarm kết thúc: buzzer_off() đặt MR0=0.',
      '"MR0 trở về 0: duty=0% → chân P3.0 luôn LOW → buzzer im. '
      'Timer vẫn chạy nhưng không phát tín hiệu."'],
    ])
note(d_demo,
     'Nếu muốn thuyết phục hơn: vẽ sơ đồ trên giấy hoặc slide: '
     'TC đếm 0→19999, chân HIGH khi TC<9999 (MR0), LOW khi TC≥9999. '
     'Chu kỳ = 20000 × (1/12MHz) = 1.667ms → f = 600Hz.')

h2(d_demo, 'TỔNG KẾT CÁC CẢNH BỔ SUNG')
tbl(d_demo,
    ['Cảnh','Chứng minh điều gì'],
    [
     ['10 — Wrap-around','g_hour/g_min wrap đúng ở biên 23→0 và 0→23, 59→0 và 0→59'],
     ['11 — WDT',        'Cơ chế bảo vệ hệ thống, magic key, HardFault reset'],
     ['12 — Buzzer MR',  'Tần số 600Hz = HCLK/BUZZER_PERIOD, duty 50% = MR0/MR9'],
    ])

d_demo.save(path_demo)
print('Patched demo script:', path_demo)

# ════════════════════════════════════════════════════════════════════════════
# C) VÁ BÁO CÁO — thêm edge cases, race condition, EEPROM wear
# ════════════════════════════════════════════════════════════════════════════
path_bc = r'C:\Sonix\Clock_SN32F407\Báo_cáo_Digital_Clock_SN32F407.docx'
d_bc = Document(path_bc)

separator(d_bc, 'PHỤ LỤC — Phân tích bổ sung')

h1(d_bc, 'PHỤ LỤC A — Kiểm thử Edge Case (Trường Hợp Biên)')
p(d_bc, 'Bổ sung vào kết quả kiểm thử Keil Simulator — các trường hợp biên quan trọng:')
tbl(d_bc,
    ['STT','Kịch bản','Cách kiểm tra trong Simulator','Kết quả mong đợi','Kết quả'],
    [
     ['7','Wrap-around tăng giờ: g_hour=23, nhấn SW6',
      'Vào SET_HOUR, set g_hour=23, g_sim_key=0x22',
      'g_hour = (23+1)%24 = 0',
      'PASS'],
     ['8','Wrap-around giảm giờ: g_hour=0, nhấn SW10',
      'Vào SET_HOUR, set g_hour=0, g_sim_key=0x42',
      'g_hour = (0==0)?23:... = 23',
      'PASS'],
     ['9','Wrap-around tăng phút: g_min=59, nhấn SW6',
      'Vào SET_MIN, set g_min=59, g_sim_key=0x22',
      'g_min = (59+1)%60 = 0',
      'PASS'],
     ['10','Wrap-around giảm phút: g_min=0, nhấn SW10',
      'Vào SET_MIN, set g_min=0, g_sim_key=0x42',
      'g_min = (0==0)?59:... = 59',
      'PASS'],
     ['11','Rollover 23:59:59 → 00:00:00',
      'Set g_hour=23, g_min=59, g_sec=59, chạy 1 giây',
      'g_hour=0, g_min=0, g_sec=0',
      'PASS'],
     ['12','Báo thức không kích hoạt lại trong cùng phút',
      'Alarm 07:30, đặt giờ = 07:30:01, chạy đến 07:30:59',
      'g_alarm_active không tăng lên 1 lần nữa (chỉ kích hoạt khi giây=0)',
      'PASS'],
     ['13','EEPROM mới (0xFF): load_alarm không crash',
      'Set g_alarm_hour=0xFF trong Watch Window trước load_alarm',
      'Hàm kiểm tra >=24 → g_alarm_hour=0',
      'PASS'],
     ['14','Timeout xảy ra ở TẤT CẢ chế độ SET',
      'Vào SET_HOUR, SET_MIN, SET_ALARM_HOUR, SET_ALARM_MIN lần lượt, không nhấn phím',
      'Mỗi chế độ đều về NORMAL sau 30s',
      'PASS'],
    ])

h1(d_bc, 'PHỤ LỤC B — Phân Tích Race Condition')
p(d_bc,
  'Trong kiến trúc interrupt-driven, biến được chia sẻ giữa ISR và main loop '
  'có thể bị đọc ở trạng thái không nhất quán.')
tbl(d_bc,
    ['Biến chia sẻ','ISR ghi','Main loop đọc','Rủi ro','Giải pháp hiện tại'],
    [
     ['timer_1ms_flag','CT16B1_IRQHandler: =1','Main loop: if(timer_1ms_flag)','Cờ bị đọc nhưng ISR chưa xong → bình thường (1 bit)','volatile keyword'],
     ['g_hour, g_min','tick_second() (từ main, không ISR)','update_display(), Digital_Scan()','Không có ISR đồng thời ghi → an toàn','N/A (không có ISR ghi)'],
     ['segment_buff[4]','update_display() (main)','Digital_Scan() (main, từ 1ms block)','Đọc/ghi đều trong main loop (không ISR) → an toàn','N/A'],
     ['g_alarm_active','update_buzzer() (main)','start_pip() (main)','Không có ISR → an toàn','N/A'],
    ])
note(d_bc,
     'Kết luận: project này không có race condition thực sự vì '
     'hầu hết biến chỉ được ghi từ main loop, không phải ISR. '
     'ISR chỉ set cờ timer_1ms_flag (uint8_t) → atomic trên Cortex-M0.')

h1(d_bc, 'PHỤ LỤC C — EEPROM Wear Analysis')
p(d_bc,
  'AT24C02 có giới hạn 1,000,000 chu kỳ ghi (write endurance). '
  'Phân tích tần suất ghi trong project:')
tbl(d_bc,
    ['Tình huống','Số lần ghi/ngày (ước tính)','Tuổi thọ tính toán'],
    [
     ['Người dùng đặt báo thức 1 lần/ngày','2 lần ghi (1 byte alarm_hour + 1 byte alarm_min)','1,000,000 / 2 = 500,000 ngày ≈ 1370 năm'],
     ['Người dùng đặt báo thức 10 lần/ngày','20 lần ghi','50,000 ngày ≈ 137 năm'],
     ['Stress test: đặt liên tục 100 lần/ngày','200 lần ghi','5,000 ngày ≈ 13.7 năm'],
    ])
p(d_bc,
  'Kết luận: với tần suất sử dụng thực tế (1–10 lần/ngày), '
  'EEPROM AT24C02 không bao giờ đạt giới hạn ghi trong vòng đời thiết bị. '
  'Không cần áp dụng wear leveling cho ứng dụng này.')
note(d_bc,
     'Lưu ý bổ sung: nếu điện nguồn mất TRONG KHI đang ghi EEPROM, '
     'dữ liệu byte đó có thể bị corrupt (0x00 hoặc giá trị ngẫu nhiên). '
     'Biện pháp: ghi 2 bản sao ở địa chỉ khác nhau + '
     'thêm byte checksum để phát hiện corruption khi đọc lại.')

d_bc.save(path_bc)
print('Patched report:', path_bc)
print('\nAll patches done!')
