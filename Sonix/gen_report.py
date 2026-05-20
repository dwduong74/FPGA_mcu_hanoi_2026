# -*- coding: utf-8 -*-
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

import docx
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = docx.Document()

for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.0)

def set_font(run, size=12, bold=False, color=None):
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.name = 'Times New Roman'
    return p

def add_para(doc, text, bold=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph(style='Normal')
    p.alignment = align
    run = p.add_run(text)
    set_font(run, 12, bold)
    return p

def add_bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    set_font(run, 12)
    return p

def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    hrow = table.rows[0]
    for i, h in enumerate(headers):
        cell = hrow.cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for ri, row in enumerate(rows):
        tr = table.rows[ri + 1]
        for ci, val in enumerate(row):
            cell = tr.cells[ci]
            cell.text = str(val)
            for run in cell.paragraphs[0].runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(11)
    return table

# ============================================================
# TRANG BÌA
# ============================================================
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Báo cáo bài thi Hệ thống nhúng – Board SN32F407_EVK')
set_font(run, 14, True)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Đề bài: Đồng hồ kỹ thuật số có báo thức trên vi điều khiển SN32F407')
set_font(run, 13, True)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Đội Thi : Sunflower (10 – PTIT)')
set_font(run, 12, True)

doc.add_paragraph()

add_para(doc, 'Danh sách thành viên đội thi :', bold=True, align=WD_ALIGN_PARAGRAPH.LEFT)
add_table(doc,
    ['Họ và tên', 'MSSV', 'Email', 'SĐT', 'Vai trò'],
    [
        ['Hoàng Bá Tước',          'B23DCDT283', 'tuochb.b23dt283@stu.ptit.edu.vn',    '0349121116', 'Trưởng Nhóm'],
        ['Nguyễn Phúc Đăng Dương', 'B23DCDK038', 'duongnpd.b23dk038@stu.ptit.edu.vn', '0967385482', 'Thành viên'],
        ['Nguyễn Anh Tú',          'B25DCDT224', 'tuna.b25dt224@stu.ptit.edu.vn',      '0846218186', 'Thành viên'],
        ['Nguyễn Đình Hồng Phúc',  'B23DCDT191', 'phucndh.b23dt191@stu.ptit.edu.vn',  '0967809536', 'Thành viên'],
    ]
)
doc.add_paragraph()

# ============================================================
# ĐỀ THI
# ============================================================
add_heading(doc, 'Đề thi', level=1)
add_para(doc,
    'Thiết kế chương trình nhúng trên board SN32F407_EVK (vi điều khiển SN32F407, '
    'ARM Cortex-M0, HCLK = 12 MHz) thực hiện các chức năng sau:'
)
for r in [
    'Hiển thị giờ và phút dạng HH.MM trên 4 LED 7 đoạn (quét động multiplexing, tần số làm tươi 250 Hz).',
    'LED D1 (P3.9) nhấp nháy 1 Hz để chỉ thị giây đang chạy (không hiển thị giây trên 7-SEG).',
    'Sử dụng timer CT16B1 tạo ngắt 1 ms, đồng bộ toàn bộ tác vụ hệ thống (không dùng delay()).',
    'Bàn phím ma trận 4×4: SW1-3 = 7-9, SW5-7 = 4-6, SW9-11 = 1-3, SW14 = 0, SW13 = (−), SW15 = (+), SW4 = Cài giờ, SW8 = Cài báo thức, SW12 = Bật/Tắt báo thức, SW16 = Tắt màn hình.',
    'Cài đặt giờ và báo thức theo Cursor Mode: con trỏ nhấp nháy từng chữ số, nhập số trực tiếp 0–9 hoặc dùng +/− có wrap-around.',
    'LED D0 (P3.8, active-LOW) nhấp nháy theo cursor khi đang ở chế độ cài báo thức.',
    'Buzzer kêu ON/OFF 500 ms khi đến giờ báo thức, tối đa 60 giây. Bất kỳ phím nào tắt tiếng ngay.',
    'SW12: bật báo thức → 2 tiếng bíp xác nhận; tắt báo thức → 1 tiếng bíp xác nhận.',
    'Lưu giờ báo thức và trạng thái bật/tắt vào EEPROM AT24C02 qua I2C; đọc lại khi khởi động.',
    'Khởi động: nhấp nháy giờ báo thức đã lưu 2 lần trên LED 7 đoạn trước khi chạy đồng hồ.',
    'SW16: tắt toàn bộ LED tiết kiệm điện. Bất kỳ phím nào bật lại.',
    'WDT (Watchdog Timer) ~250 ms bảo vệ hệ thống khỏi treo chương trình.',
    'Sau 30 giây không thao tác ở chế độ cài đặt, tự thoát về NORMAL.',
]:
    add_bullet(doc, r)

# ============================================================
# CHƯƠNG I
# ============================================================
add_heading(doc, 'Chương I – Báo cáo nội dung đã thực hiện', level=1)
add_heading(doc, '1. Tính năng đã thực hiện trong bài thi', level=2)

for title, desc in [
    (
        'a. Hiển thị HH.MM trên LED 7 đoạn (quét động)',
        'Bốn LED 7 đoạn được quét lần lượt theo phương pháp multiplexing: mỗi 1 ms kích hoạt một COM '
        '(COM0–COM3 qua transistor NPN, P1.9–P1.12), đồng thời xuất dữ liệu tương ứng qua 8 chân segment '
        '(P0.0–P0.7, active-HIGH). SEGMENT_TABLE[] trong Segment.c chứa mã 7 đoạn cho chữ số 0–9. '
        'Bit DP (0x80) được OR vào COM1 (chữ số đơn vị giờ) để hiển thị dấu chấm phân cách. '
        'Tần số quét: 4 COM × 1 ms = 250 Hz — mắt người không nhận ra nhấp nháy.'
    ),
    (
        'b. Timer CT16B1 – khung thời gian 1 ms',
        'CT16B1 cấu hình chế độ Match Reset với MR9 = 11999 tại HCLK = 12 MHz, tạo ngắt chính xác mỗi 1 ms. '
        'ISR chỉ đặt cờ timer_1ms_flag = 1. Vòng while(1) kiểm tra cờ này và thực hiện tuần tự 5 tác vụ: '
        'quét LED, quét phím, cập nhật giây/phút/giờ, điều khiển LED D1, và cập nhật buzzer. '
        'Thiết kế này loại bỏ hoàn toàn hàm delay() và đảm bảo mọi ngoại vi được phục vụ đúng chu kỳ.'
    ),
    (
        'c. State machine 4 trạng thái',
        'STATE_NORMAL (0): đồng hồ chạy bình thường, LED D1 nhấp nháy 1 Hz. '
        'STATE_SET_HOUR (1): cài giờ hiện tại theo Cursor Mode, con trỏ nhấp nháy tại chữ số đang chỉnh. '
        'STATE_SET_ALARM (2): cài báo thức tương tự, LED D0 nhấp nháy theo cursor. '
        'STATE_DISPLAY_OFF (3): tắt toàn bộ LED 7-SEG, D0, D1 tiết kiệm điện.'
    ),
    (
        'd. Cursor Mode – nhập số trực tiếp',
        'Biến g_cursor (0–3) trỏ vào 4 chữ số: H_tens, H_units, M_tens, M_units. '
        'Chữ số tại con trỏ nhấp nháy 500 ms ON / 500 ms OFF. '
        'Nhấn phím số 0–9: hàm set_cursor_digit() tính lại giờ/phút và clamp nếu vượt giới hạn '
        '(VD: nhấn 9 ở H_tens → giờ clamp về 23). '
        'Nhấn (+)/(−): gọi inc_cursor()/dec_cursor() với wrap-around: 23→0, 0→23, 59→0, 0→59.'
    ),
    (
        'e. LED D1 nhấp nháy chỉ thị giây',
        'Vì 7-SEG chỉ hiển thị HH.MM, giây không thể hiện trực tiếp. '
        'LED D1 (P3.9, active-LOW) bật khi g_ms_cnt < 500 và tắt khi g_ms_cnt ≥ 500 '
        'trong mỗi chu kỳ 1000 ms, tạo nhấp nháy đúng 1 Hz. '
        'D1 tắt hoàn toàn khi ở STATE_DISPLAY_OFF.'
    ),
    (
        'f. Startup Animation – hiển thị giờ báo thức khi bật nguồn',
        'Ngay sau khi đọc EEPROM, firmware chạy vòng blocking riêng: nhấp nháy giờ báo thức đã lưu '
        'ON 400 ms → OFF 400 ms, lặp 2 lần (tổng 1,6 giây). '
        'Trong vòng này vẫn quét LED và feed WDT. '
        'Mục đích: cho phép BTC/người dùng xác nhận giờ báo thức đã lưu mà không cần vào menu. '
        'Sau animation, đồng hồ bắt đầu từ 00:00.'
    ),
    (
        'g. Báo thức – buzzer ON/OFF 500 ms, tối đa 60 giây',
        'Khi g_hour == g_alarm_hour && g_min == g_alarm_min (tại thời điểm phút mới bắt đầu), '
        'start_alarm_beep() kích hoạt: buzzer kêu nhịp nhàng ON 500 ms / OFF 500 ms, '
        'biến g_alarm_remain_ms đếm ngược 60 000 ms. Hết thời gian → tự tắt. '
        'Bất kỳ phím nào nhấn trong khi báo thức đang kêu → dừng ngay, nuốt phím.'
    ),
    (
        'h. SW12 – bật/tắt báo thức với phản hồi âm thanh',
        'Nhấn SW12 khi báo thức đang BẬT: tắt báo thức (g_alarm_enabled = 0) + phát 1 tiếng bíp. '
        'Nhấn SW12 khi báo thức đang TẮT: bật báo thức (g_alarm_enabled = 1) + phát 2 tiếng bíp. '
        'Hàm start_pips(n) quản lý chuỗi bíp qua biến g_pip_remain và g_pip_gap_cnt '
        'trong update_buzzer() — không dùng delay(). '
        'Trạng thái mới được lưu vào EEPROM ngay lập tức.'
    ),
    (
        'i. Lưu/đọc báo thức vào EEPROM AT24C02 qua I2C (Smart Write)',
        'EEPROM layout: địa chỉ 0x00 = alarm_hour, 0x01 = alarm_min, '
        '0x02 = trạng thái bật/tắt (0x01/0x00). '
        'Smart Write Strategy: chỉ ghi EEPROM khi người dùng xác nhận xong cài báo thức '
        '(SW8 lần thứ 4) hoặc SW12 toggle. '
        'Giờ hiện tại KHÔNG lưu vào EEPROM để tránh wear. '
        'AT24C02 chịu được 1 000 000 chu kỳ ghi; với tần suất thực tế (1–10 lần/ngày), '
        'EEPROM dùng được 137–1 370 năm.'
    ),
    (
        'j. WDT – Watchdog Timer ~250 ms',
        'WDT được khởi tạo SAU tất cả thao tác I2C để tránh reset sớm trong quá trình init. '
        'Lệnh __WDT_FEED_VALUE đặt tại đầu while(1), đảm bảo feed mỗi chu kỳ 1 ms. '
        'Nếu firmware treo (I2C hang, stack overflow…), WDT tự reset MCU sau ~250 ms.'
    ),
    (
        'k. Debounce ma trận phím',
        'KeyScan() chạy mỗi 1 ms với thuật toán 2 chiều (Row/Column): '
        'set ROW output LOW đọc COL → phát hiện cột; '
        'đổi chiều set COL output LOW đọc ROW → xác định hàng. '
        'Debounce 50 ms: phím phải ổn định liên tục 50 lần kiểm tra mới ghi nhận. '
        'Sau 200 ms giữ phím liên tục, phím bị hủy tránh nhấn lặp.'
    ),
]:
    add_heading(doc, title, level=3)
    add_bullet(doc, desc)

# -------- Kiến trúc module --------
add_heading(doc, '2. Kiến trúc module bài thi', level=2)
add_para(doc, 'Sơ đồ phân cấp module:')
add_para(doc,
    'main.c  (UserAPP)\n'
    '├── CT16B1.c         – Timer ngắt 1 ms\n'
    '├── CT16B0.c         – Buzzer PWM 600 Hz\n'
    '├── Segment.c        – Mã hóa 7 đoạn, quét LED 4 COM\n'
    '├── KeyScan.c        – Ma trận phím 4×4, debounce 50 ms\n'
    '├── I2C0.c           – Giao tiếp I2C master (polling)\n'
    '├── EEPROM.c         – eeprom_read(), eeprom_write()\n'
    '├── GPIO.c           – Cấu hình chân I/O\n'
    '├── PFPA.c           – Pin Function Assignment (I2C0→P0.10/11)\n'
    '└── WDT.c            – Watchdog Timer ~250 ms'
)

add_heading(doc, 'Bảng mô tả module', level=3)
add_table(doc,
    ['Module', 'File', 'Vai trò'],
    [
        ['main.c',  'UserAPP/main.c',   'Vòng lặp chính, state machine, logic đồng hồ & báo thức'],
        ['CT16B1',  'Driver/CT16B1.c',  'Tạo ngắt 1 ms; ISR đặt timer_1ms_flag'],
        ['CT16B0',  'Driver/CT16B0.c',  'PWM buzzer 600 Hz: MR9 = 19999, MR0 = 9999 (duty 50%)'],
        ['Segment', 'Module/Segment.c', 'SEGMENT_TABLE[], Digital_Scan() quét 4 COM mỗi 1 ms'],
        ['KeyScan', 'Module/KeyScan.c', 'Ma trận phím 4×4, debounce 50 ms, trả về mã KEY_x'],
        ['EEPROM',  'Module/EEPROM.c',  'eeprom_write(), eeprom_read() đóng gói I2C AT24C02'],
        ['I2C0',    'Driver/I2C0.c',    'I2C master polling: START / STOP / WRITE / READ'],
        ['GPIO',    'Driver/GPIO.c',    'GPIO_Init(): SEG P0, COM P1.9–12, KEY P1.4–7 + P2.4–7, LED P3.8–9'],
        ['PFPA',    'Driver/PFPA.c',    'PFPA_Init(): I2C0→P0.10/11, CT16B0 PWM0→P3.0'],
        ['WDT',     'Driver/WDT.c',     'WDT_Init(), __WDT_FEED_VALUE, timeout ~250 ms'],
    ]
)

add_heading(doc, 'Bảng chuyển trạng thái', level=3)
add_table(doc,
    ['Từ trạng thái', 'Đến trạng thái', 'Điều kiện'],
    [
        ['NORMAL',      'SET_HOUR',    'Nhấn SW4'],
        ['NORMAL',      'SET_ALARM',   'Nhấn SW8'],
        ['NORMAL',      'DISPLAY_OFF', 'Nhấn SW16'],
        ['SET_HOUR',    'SET_HOUR',    'Nhấn SW4 → cursor tiến: 0 → 1 → 2 → 3'],
        ['SET_HOUR',    'NORMAL',      'SW4 sau cursor 3 (g_sec = 0) hoặc Timeout 30 s'],
        ['SET_HOUR',    'SET_HOUR',    'Phím 0–9: set_cursor_digit(); Phím +/−: inc/dec_cursor()'],
        ['SET_ALARM',   'SET_ALARM',   'Nhấn SW8 → cursor tiến: 0 → 1 → 2 → 3'],
        ['SET_ALARM',   'NORMAL',      'SW8 sau cursor 3 → lưu EEPROM; hoặc Timeout 30 s'],
        ['SET_ALARM',   'SET_ALARM',   'Phím 0–9: set_cursor_digit(); Phím +/−: inc/dec_cursor()'],
        ['DISPLAY_OFF', 'NORMAL',      'Bất kỳ phím nào'],
        ['NORMAL',      'NORMAL',      'SW12: toggle g_alarm_enabled, lưu EEPROM, phát 1 hoặc 2 bíp'],
    ]
)

# -------- Flow hoạt động --------
add_heading(doc, '3. Flow hoạt động', level=2)
add_heading(doc, 'a) Sơ đồ khối chương trình', level=3)
add_para(doc,
    'Chương trình khởi động theo trình tự: SystemInit → PFPA_Init → GPIO_Init → '
    'CT16B0_Init (PWM buzzer) → CT16B1_Init (timer 1 ms) → I2C0_Init → '
    'load_settings_from_eeprom() → update_display() → WDT_Init → '
    'Startup Animation (nhấp nháy giờ báo thức 2 lần) → vòng lặp while(1).'
)
add_para(doc, 'Trong vòng lặp chính (mỗi 1 ms khi timer_1ms_flag = 1):')
for s in [
    'Feed WDT (__WDT_FEED_VALUE) – đầu tiên, trước mọi xử lý.',
    'Task 1: Digital_Scan() – quét 1 COM, cập nhật LED 7 đoạn.',
    'Task 2: KeyScan() + process_key() – đọc phím, xử lý state machine.',
    'Task 3: g_ms_cnt++ → mỗi 1 000 ms gọi tick_second() cập nhật giờ và kiểm tra báo thức.',
    'Task 3b: Điều khiển LED D1 (ON khi g_ms_cnt < 500, OFF khi ≥ 500).',
    'Task 4: Nhấp nháy cursor (toggle mỗi 500 ms) + kiểm tra timeout 30 s.',
    'Task 5: update_buzzer() – quản lý chuỗi bíp phản hồi và báo thức.',
]:
    add_bullet(doc, s)

add_heading(doc, 'b) Kết quả kiểm thử', level=3)
add_table(doc,
    ['STT', 'Kịch bản kiểm thử', 'Nội dung chi tiết', 'Kết quả'],
    [
        ['1',  'Đếm giờ cơ bản',           'Chạy Free-Run, g_sec tăng 0→59, g_min tăng, g_hour tăng, rollover 23:59:59 → 00:00:00.', 'PASS'],
        ['2',  'Cài giờ Cursor Mode',       'Nhấn SW4 vào SET_HOUR. Nhấn số 1→H_tens=1, 4→H_units=4, 3→M_tens=3, 0→M_units=0. Màn hình hiển thị 14:30. Nhấn SW4 lần 4 thoát về NORMAL.', 'PASS'],
        ['3',  'Phím +/− wrap-around',      'SET_HOUR: g_hour=23, nhấn SW15 → g_hour=0. g_min=0, nhấn SW13 → g_min=59.', 'PASS'],
        ['4',  'Clamp giá trị',             'Cursor=0, nhấn số 9 → giờ clamp về 23 (không phải 90). Cursor=2, nhấn 7 → phút clamp về 59.', 'PASS'],
        ['5',  'Cài báo thức & lưu EEPROM','Nhấn SW8, nhập 4 chữ số báo thức, SW8 lần cuối → lưu EEPROM. Rút điện cắm lại → startup animation hiện đúng giờ báo thức vừa cài.', 'PASS'],
        ['6',  'Startup animation',         'Sau khi cắm nguồn: LED nhấp nháy giờ báo thức ON 400 ms / OFF 400 ms đúng 2 lần trước khi chạy từ 00:00.', 'PASS'],
        ['7',  'SW12 toggle báo thức',      'Nhấn SW12 khi BẬT → 1 bíp. Nhấn SW12 khi TẮT → 2 bíp.', 'PASS'],
        ['8',  'Báo thức kêu 60 giây',      'Cài báo thức = giờ hiện tại + 1 phút. Khi đến giờ: buzzer ON/OFF 500 ms, tự tắt sau 60 s.', 'PASS'],
        ['9',  'Bất kỳ phím tắt báo thức', 'Trong khi báo thức kêu, nhấn SW1 → buzzer dừng ngay, phím bị nuốt.', 'PASS'],
        ['10', 'Timeout 30 giây',           'Vào SET_HOUR, không nhấn phím 30 s → tự về NORMAL kèm 1 bíp.', 'PASS'],
        ['11', 'Tắt/bật màn hình SW16',     'Nhấn SW16 → toàn bộ LED 7-SEG, D0, D1 tắt. Nhấn SW1 → bật lại, giờ hiển thị đúng.', 'PASS'],
        ['12', 'LED D1 chỉ thị giây',       'D1 sáng 500 ms / tắt 500 ms đúng 1 Hz trong trạng thái NORMAL và SET.', 'PASS'],
        ['13', 'Rollover không hiện XX:60', 'g_min tăng đến 60 → tick_second() wrap về 0 TRƯỚC khi gọi update_display() – màn hình không bao giờ hiển thị :60.', 'PASS'],
        ['14', 'WDT không reset bình thường','Firmware chạy liên tục 10 phút, __WDT_FEED_VALUE trong while(1) đảm bảo không reset.', 'PASS'],
    ]
)

add_heading(doc, 'c) Nhận xét kết quả', level=3)
for c in [
    'State machine chuyển trạng thái đúng theo mọi nhánh điều kiện.',
    'Cursor Mode hoạt động mượt mà: nhập số trực tiếp clamp đúng, +/− wrap-around chính xác.',
    'segment_buff[1] luôn có bit DP (0x80) để hiển thị dấu chấm phân cách HH.MM.',
    'Bug hiển thị 23:60 đã được sửa: update_display() chỉ gọi SAU khi toàn bộ rollover phút/giờ hoàn tất trong tick_second().',
    'Smart Write Strategy đảm bảo EEPROM không bị ghi quá mức – tuổi thọ hàng trăm năm với tần suất sử dụng thực tế.',
    'Startup animation là điểm nổi bật giúp BTC xác nhận EEPROM hoạt động đúng ngay khi bật nguồn.',
]:
    add_bullet(doc, c)

# -------- Phần cứng --------
add_heading(doc, '4. Mô tả phần cứng sử dụng', level=2)
add_table(doc,
    ['Tín hiệu', 'Chân MCU', 'Phần cứng', 'Ghi chú'],
    [
        ['SEG A–G, DP', 'P0.0–P0.7',  'LED 7 đoạn segment',     'GPIO output, active-HIGH'],
        ['COM0–COM3',   'P1.9–P1.12', 'Anode chung (mux scan)',  'GPIO output qua transistor NPN'],
        ['Key ROW 1–4', 'P1.4–P1.7', 'Hàng ma trận phím',      'GPIO output'],
        ['Key COL 1–4', 'P2.4–P2.7', 'Cột ma trận phím',       'GPIO input pull-up'],
        ['LED D0',      'P3.8',       'LED báo trạng thái alarm','Active-LOW (BCLR = sáng)'],
        ['LED D1',      'P3.9',       'LED nhịp giây',           'Active-LOW, nhấp nháy 1 Hz'],
        ['Buzzer',      'P3.0',       'Buzzer passive',          'CT16B0 PWM0, 600 Hz, duty 50%'],
        ['I2C0 SCL/SDA','P0.10/P0.11','EEPROM AT24C02',         'PFPA option 2, I2C polling'],
        ['HCLK',        'IHRC',       '12 MHz nội',              'Không cần thạch anh ngoài'],
    ]
)

# -------- Timing --------
add_heading(doc, '5. Bảng timing tham khảo', level=2)
add_table(doc,
    ['Tham số', 'Giá trị', 'Ý nghĩa'],
    [
        ['HCLK',             '12 MHz',          '~83 ns/cycle'],
        ['CT16B1 MR9',       '11 999',          'Ngắt timer chính xác 1 ms'],
        ['LED scan period',  '1 ms/COM',        '4 COM × 1 ms = 250 Hz refresh (không flicker)'],
        ['Blink half-period','500 ms',          'Cursor nhấp nháy khi cài đặt'],
        ['Timeout SET mode', '30 s = 30 000 ms','Tự về NORMAL nếu không nhấn phím'],
        ['Pip duration',     '200 ms',          'Mỗi tiếng bíp phản hồi phím'],
        ['Pip gap',          '200 ms',          'Khoảng lặng giữa 2 bíp (cho SW12 BẬT)'],
        ['Alarm duration',   '60 s = 60 000 ms','Buzzer kêu tối đa 60 giây'],
        ['Alarm ON/OFF',     '500 ms / 500 ms', 'Nhịp nhàng báo thức'],
        ['Startup animation','4 × 400 ms',      '2 lần nhấp nháy giờ báo thức khi bật nguồn'],
        ['WDT timeout',      '~250 ms',         'Reset MCU nếu không feed'],
        ['CT16B0 MR9',       '19 999',          'Chu kỳ PWM buzzer: 12 MHz / 20 000 = 600 Hz'],
        ['CT16B0 MR0',       '9 999',           'Duty cycle 50%'],
        ['I2C speed',        '100 kHz',         'Standard mode, AT24C02'],
        ['Debounce',         '50 ms',           'Phím phải ổn định 50 lần kiểm tra liên tiếp'],
    ]
)

# ============================================================
# CHƯƠNG II
# ============================================================
add_heading(doc, 'Chương II – Hướng dẫn build – nạp – chạy demo', level=1)

add_heading(doc, '1. Phần cứng – phần mềm sử dụng', level=2)
add_para(doc, 'a) Phần cứng', bold=True, align=WD_ALIGN_PARAGRAPH.LEFT)
for h in [
    'Board SN32F407_EVK (vi điều khiển SN32F407, ARM Cortex-M0, HCLK = 12 MHz).',
    '4 LED 7 đoạn cathode chung trên board (COM khuếch đại qua transistor NPN).',
    'Bàn phím ma trận 4×4 (16 phím SW1–SW16) trên board.',
    'LED D0 (P3.8) và LED D1 (P3.9) trên board.',
    'Buzzer passive kết nối CT16B0 PWM (P3.0).',
    'EEPROM AT24C02 trên board (I2C0 – P0.10 SCL, P0.11 SDA).',
    'Cáp nạp CMSIS-DAP hoặc SN-Link USB Downloader.',
]:
    add_bullet(doc, h)

add_para(doc, 'b) Phần mềm', bold=True, align=WD_ALIGN_PARAGRAPH.LEFT)
for s in [
    'Keil MDK-ARM v5 (uVision5).',
    'ARM Compiler 6 (armclang V6.24), cài tại C:\\Keil_v5\\ARM\\ARMCLANG.',
    'SONiX SN32F4_DFP Pack v1.1.1 (cài trong Keil Pack Manager).',
    'Sonix ISP Tool hoặc Keil Flash Download để nạp firmware.',
]:
    add_bullet(doc, s)

add_heading(doc, '2. Build chương trình', level=2)
add_para(doc, '2.1 Mở project trong Keil uVision5', bold=True, align=WD_ALIGN_PARAGRAPH.LEFT)
for b in [
    'Mở file Clock_SN32F407\\Clock_SN32F407.uvprojx.',
    'Kiểm tra Device: SN32F407F; Compiler: Use default compiler version 6.',
    'Include paths: .\\Source\\Driver; .\\Source\\Module.',
    'Đảm bảo ADC.c KHÔNG có trong project (đã loại bỏ).',
]:
    add_bullet(doc, b)

add_para(doc, '2.2 Build project', bold=True, align=WD_ALIGN_PARAGRAPH.LEFT)
for b in [
    'Nhấn F7 (Build) hoặc Project → Build Target.',
    'Kết quả mong đợi: "0 Error(s), 0 Warning(s)".',
    'Output: Clock_SN32F407\\Objects\\Clock_SN32F407.axf (và .hex).',
    'Lưu ý: CT16B0_IRQHandler không dùng từ khóa __irq (không hợp lệ trên ARM Compiler 6 / armclang). Handler được nhận diện tự động qua vector table Cortex-M.',
]:
    add_bullet(doc, b)

add_heading(doc, '3. Nạp firmware lên board SN32F407_EVK', level=2)
add_para(doc, '3.1 Cấu hình Flash Download trong Keil', bold=True, align=WD_ALIGN_PARAGRAPH.LEFT)
add_bullet(doc, 'Alt+F7 → tab Debug → chọn "Use: Sonix SN-Link Cortex-M Debugger" (hoặc CMSIS-DAP).')
add_bullet(doc, 'Tab Flash Download: chọn đúng algorithm SN32F407.')

add_para(doc, '3.2 Nạp chương trình', bold=True, align=WD_ALIGN_PARAGRAPH.LEFT)
add_bullet(doc, 'Kết nối board SN32F407_EVK qua cáp USB CMSIS-DAP.')
add_bullet(doc, 'Nhấn F8 (Download) trong Keil hoặc dùng Sonix ISP Tool với file .hex.')
add_bullet(doc, 'Kết quả: "Programming Done", "Verify OK".')

add_heading(doc, '4. Chạy chương trình và kết quả demo', level=2)
for i, s in enumerate([
    'Cắm nguồn vào board. LED 7 đoạn nhấp nháy giờ báo thức đã lưu 2 lần (startup animation).',
    'Đồng hồ bắt đầu chạy từ 00.00. LED D1 nhấp nháy 1 Hz.',
    'Nhấn SW4 vào SET_HOUR. Chữ số H_tens nhấp nháy. Nhập số bằng numpad (SW1-3=7-9, SW5-7=4-6, SW9-11=1-3, SW14=0). Dùng SW15/SW13 để tăng/giảm.',
    'Nhấn SW4 mỗi lần để tiến cursor (H_tens→H_units→M_tens→M_units). Sau ô cuối tự về NORMAL.',
    'Nhấn SW8 vào SET_ALARM. LED D0 nhấp nháy cùng cursor. Nhập 4 chữ số. SW8 lần cuối lưu EEPROM.',
    'Nhấn SW12 để tắt báo thức → 1 bíp. Nhấn lại để bật báo thức → 2 bíp.',
    'Khi đến giờ báo thức: buzzer kêu ON/OFF 500 ms. Nhấn bất kỳ phím để tắt ngay.',
    'Nhấn SW16 → toàn bộ LED tắt. Nhấn bất kỳ phím → bật lại đúng giờ.',
    'Rút điện, cắm lại → startup animation hiện đúng giờ báo thức đã cài → xác nhận EEPROM hoạt động.',
], 1):
    add_bullet(doc, f'{i}. {s}')

# ============================================================
# PHỤ LỤC
# ============================================================
add_heading(doc, 'Phụ lục A – Phân tích EEPROM Wear', level=1)
add_para(doc,
    'AT24C02 có giới hạn 1 000 000 chu kỳ ghi. '
    'Phân tích tần suất ghi trong project với Smart Write Strategy:'
)
add_table(doc,
    ['Tình huống', 'Số lần ghi/ngày', 'Tuổi thọ tính toán'],
    [
        ['Đặt báo thức 1 lần/ngày',  '3 lần (hour + min + enabled)', '333 333 ngày ≈ 913 năm'],
        ['Đặt báo thức 10 lần/ngày', '30 lần',                       '33 333 ngày ≈ 91 năm'],
        ['Toggle SW12 100 lần/ngày', '100 lần (1 byte enabled)',      '10 000 ngày ≈ 27 năm'],
    ]
)
add_para(doc,
    'Kết luận: với tần suất sử dụng thực tế (1–10 lần/ngày), EEPROM AT24C02 không bao giờ đạt '
    'giới hạn ghi trong vòng đời thiết bị. Không cần áp dụng wear leveling cho ứng dụng này.'
)

add_heading(doc, 'Phụ lục B – Phân tích Race Condition', level=1)
add_para(doc,
    'Kết luận: project này không có race condition thực sự. '
    'ISR của CT16B1 chỉ set cờ timer_1ms_flag (uint8_t, atomic trên Cortex-M0). '
    'Mọi biến thời gian (g_hour, g_min, g_sec, segment_buff) chỉ được ghi từ main loop '
    '— không có ISR nào ghi đồng thời, không cần critical section.'
)
add_table(doc,
    ['Biến chia sẻ', 'ISR ghi', 'Main loop đọc/ghi', 'Kết luận'],
    [
        ['timer_1ms_flag', 'CT16B1_IRQHandler: = 1',  'Main: đọc & xóa',                   'Atomic 1 byte – an toàn'],
        ['g_hour, g_min',  'Không (tick_second từ main)', 'update_display(), process_key()', 'Không có ISR ghi – an toàn'],
        ['segment_buff[4]','Không',                   'update_display() & Digital_Scan()',   'Cùng luồng main – an toàn'],
        ['g_alarm_active', 'Không',                   'update_buzzer() & process_key()',     'Cùng luồng main – an toàn'],
    ]
)

out_path = r'C:\Sonix\Clock_SN32F407\BAO_CAO_CUOI_CUNG.docx'
doc.save(out_path)
print(f'Saved: {out_path}')
